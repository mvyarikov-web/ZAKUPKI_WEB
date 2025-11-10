"""Тест индексации документа из blob."""
import pytest
import uuid
from io import BytesIO
from werkzeug.datastructures import FileStorage
from sqlalchemy import text

from webapp.services.blob_storage_service import BlobStorageService
from webapp.services.db_indexing import index_document_to_db
from webapp.models.rag_models import RAGDatabase


@pytest.fixture
def rag_db(app):
    """Фикстура для RAGDatabase."""
    # RAGDatabase автоматически получит database_url из конфигурации или окружения
    return RAGDatabase(None)


@pytest.fixture
def blob_service(app):
    """Фикстура для BlobStorageService."""
    from webapp.config.config_service import ConfigService
    config = ConfigService()
    return BlobStorageService(config)


def test_index_document_from_blob_txt(db, test_user, blob_service, rag_db):
    """Тест: загрузка TXT в blob → индексация → проверка chunks."""
    # 1. Загружаем файл в blob (уникальный контент для каждого запуска)
    unique_id = str(uuid.uuid4())[:8]
    content = f"Тестовый документ для индексации {unique_id}. Содержит несколько предложений. Проверка индексации из blob."
    file_bytes = content.encode('utf-8')
    
    fake_file = FileStorage(
        stream=BytesIO(file_bytes),
        filename=f"test_index_{unique_id}.txt",
        content_type="text/plain"
    )
    
    document, is_new = blob_service.save_file_to_db(
        db=db,
        file=fake_file,
        user_id=test_user.id,
        user_path=f"test/index_{unique_id}.txt"
    )
    
    assert is_new is True
    assert document.blob is not None
    sha256_hash = document.sha256
    
    # 2. Индексируем документ (функция должна прочитать из blob)
    file_info = {
        'sha256': sha256_hash,
        'size': len(file_bytes),
        'content_type': 'text/plain'
    }
    
    doc_id, indexing_cost = index_document_to_db(
        db=rag_db,
        file_path="",  # Путь пустой - должно читаться из blob
        file_info=file_info,
        user_id=test_user.id,
        original_filename="test_index.txt",
        user_path="test/index.txt",
        chunk_size_tokens=50,
        chunk_overlap_tokens=10
    )
    
    assert doc_id is not None
    assert indexing_cost >= 0
    
    # 3. Проверяем что chunks созданы
    chunks = db.execute(
        text("SELECT id, text, chunk_idx FROM chunks WHERE document_id = :doc_id ORDER BY chunk_idx"),
        {"doc_id": doc_id}
    ).fetchall()
    
    assert len(chunks) > 0, "Должен быть создан хотя бы один chunk"
    
    # Проверяем содержимое первого chunk
    first_chunk = chunks[0]
    chunk_content = first_chunk[1]
    
    assert "Тестовый документ" in chunk_content or "индексации" in chunk_content or unique_id in chunk_content, \
        f"Chunk должен содержать текст из документа, получено: {chunk_content[:100]}"


def test_index_document_from_blob_docx(db, test_user, blob_service, rag_db):
    """Тест: загрузка DOCX в blob → индексация → проверка chunks."""
    try:
        from docx import Document
        
        # Создаём DOCX
        doc = Document()
        doc.add_paragraph("Первый параграф с важной информацией.")
        doc.add_paragraph("Второй параграф для тестирования индексации.")
        doc.add_paragraph("Третий параграф содержит ключевые слова: закупки, тендер, конкурс.")
        
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        fake_file = FileStorage(
            stream=docx_buffer,
            filename="test_index.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Загружаем в blob
        document, is_new = blob_service.save_file_to_db(
            db=db,
            file=fake_file,
            user_id=test_user.id,
            user_path="test/index.docx"
        )
        
        assert is_new is True
        assert document.blob is not None
        
        # Индексируем
        file_info = {
            'sha256': document.sha256,
            'size': document.size_bytes,
            'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
        doc_id, indexing_cost = index_document_to_db(
            db=rag_db,
            file_path="",
            file_info=file_info,
            user_id=test_user.id,
            original_filename="test_index.docx",
            user_path="test/index.docx",
            chunk_size_tokens=50,
            chunk_overlap_tokens=10
        )
        
        assert doc_id is not None
        
        # Проверяем chunks
        chunks = db.execute(
            text("SELECT text FROM chunks WHERE document_id = :doc_id"),
            {"doc_id": doc_id}
        ).fetchall()
        
        assert len(chunks) > 0, "Должны быть созданы chunks"
        
        # Проверяем что хотя бы один chunk содержит контент
        all_content = " ".join([c[0] for c in chunks])
        assert "параграф" in all_content.lower() or "информацией" in all_content.lower(), \
            f"Chunks должны содержать текст из DOCX"
        
    except ImportError:
        pytest.skip("python-docx not available")


def test_index_duplicate_document_skips_reindexing(db, test_user, blob_service, rag_db):
    """Тест: повторная индексация того же документа не создаёт дубликаты chunks."""
    unique_id = str(uuid.uuid4())[:8]
    content = f"Уникальный текст для проверки дедупликации индексации {unique_id}."
    file_bytes = content.encode('utf-8')
    
    fake_file = FileStorage(
        stream=BytesIO(file_bytes),
        filename=f"duplicate_test_{unique_id}.txt",
        content_type="text/plain"
    )
    
    # Первая загрузка
    document, is_new = blob_service.save_file_to_db(
        db=db,
        file=fake_file,
        user_id=test_user.id,
        user_path=f"test/duplicate_{unique_id}.txt"
    )
    
    file_info = {
        'sha256': document.sha256,
        'size': len(file_bytes),
        'content_type': 'text/plain'
    }
    
    # Первая индексация
    doc_id_1, _ = index_document_to_db(
        db=rag_db,
        file_path="",
        file_info=file_info,
        user_id=test_user.id,
        original_filename=f"duplicate_test_{unique_id}.txt",
        user_path=f"test/duplicate_{unique_id}.txt",
        chunk_size_tokens=50
    )
    
    chunks_count_1 = db.execute(
        text("SELECT COUNT(*) FROM chunks WHERE document_id = :doc_id"),
        {"doc_id": doc_id_1}
    ).fetchone()[0]
    
    # Вторая индексация того же файла
    doc_id_2, _ = index_document_to_db(
        db=rag_db,
        file_path="",
        file_info=file_info,
        user_id=test_user.id,
        original_filename=f"duplicate_test_{unique_id}.txt",
        user_path=f"test/duplicate_{unique_id}.txt",
        chunk_size_tokens=50
    )
    
    # ID документа должен совпадать
    assert doc_id_1 == doc_id_2, "Повторная индексация должна вернуть тот же document_id"
    
    # Количество chunks не должно увеличиться
    chunks_count_2 = db.execute(
        text("SELECT COUNT(*) FROM chunks WHERE document_id = :doc_id"),
        {"doc_id": doc_id_2}
    ).fetchone()[0]
    
    assert chunks_count_1 == chunks_count_2, "Повторная индексация не должна создавать дубликаты chunks"
