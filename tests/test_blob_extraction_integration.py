"""Интеграционный тест: загрузка в blob + извлечение текста из blob."""
import pytest
from io import BytesIO
from werkzeug.datastructures import FileStorage

from webapp.services.blob_storage_service import BlobStorageService
from document_processor.extractors.text_extractor import extract_text_from_bytes


@pytest.fixture
def blob_service(config):
    """Фикстура для BlobStorageService."""
    return BlobStorageService(config)


def test_upload_txt_to_blob_and_extract(db, test_user, blob_service):
    """Тест: загрузка TXT в blob → извлечение текста из blob."""
    # Создаём тестовый TXT файл
    content = "Тестовый документ с кириллицей и числами: 12345"
    file_bytes = content.encode('utf-8')
    
    fake_file = FileStorage(
        stream=BytesIO(file_bytes),
        filename="test_document.txt",
        content_type="text/plain"
    )
    
    # Сохраняем в blob
    document, is_new = blob_service.save_file_to_db(
        db=db,
        file=fake_file,
        user_id=test_user.id,
        user_path="test/document.txt"
    )
    
    assert is_new is True
    assert document.blob is not None
    assert document.size_bytes == len(file_bytes)
    
    # Извлекаем текст из blob
    retrieved_bytes = blob_service.get_file_bytes(db, document.id)
    assert retrieved_bytes is not None
    
    extracted_text = extract_text_from_bytes(retrieved_bytes, 'txt')
    
    assert "Тестовый документ" in extracted_text
    assert "кириллицей" in extracted_text
    assert "12345" in extracted_text


def test_upload_docx_to_blob_and_extract(db, test_user, blob_service):
    """Тест: загрузка DOCX в blob → извлечение текста из blob."""
    try:
        from docx import Document
        
        # Создаём DOCX
        doc = Document()
        doc.add_paragraph("Параграф 1: Проверка кириллицы")
        doc.add_paragraph("Параграф 2: Числа и символы 123!@#")
        
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        fake_file = FileStorage(
            stream=docx_buffer,
            filename="test_document.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Сохраняем в blob
        document, is_new = blob_service.save_file_to_db(
            db=db,
            file=fake_file,
            user_id=test_user.id,
            user_path="test/document.docx"
        )
        
        assert is_new is True
        assert document.blob is not None
        
        # Извлекаем текст из blob
        retrieved_bytes = blob_service.get_file_bytes(db, document.id)
        extracted_text = extract_text_from_bytes(retrieved_bytes, 'docx')
        
        assert "Параграф 1" in extracted_text
        assert "кириллицы" in extracted_text
        assert "Параграф 2" in extracted_text
        assert "123" in extracted_text
        
    except ImportError:
        pytest.skip("python-docx not available")


def test_upload_xlsx_to_blob_and_extract(db, test_user, blob_service):
    """Тест: загрузка XLSX в blob → извлечение текста из blob."""
    try:
        from openpyxl import Workbook
        
        # Создаём XLSX
        wb = Workbook()
        ws = wb.active
        ws['A1'] = "Столбец A"
        ws['B1'] = "Столбец B"
        ws['A2'] = "Кириллица"
        ws['B2'] = 999
        
        xlsx_buffer = BytesIO()
        wb.save(xlsx_buffer)
        xlsx_buffer.seek(0)
        
        fake_file = FileStorage(
            stream=xlsx_buffer,
            filename="test_spreadsheet.xlsx",
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
        # Сохраняем в blob
        document, is_new = blob_service.save_file_to_db(
            db=db,
            file=fake_file,
            user_id=test_user.id,
            user_path="test/spreadsheet.xlsx"
        )
        
        assert is_new is True
        assert document.blob is not None
        
        # Извлекаем текст из blob
        retrieved_bytes = blob_service.get_file_bytes(db, document.id)
        extracted_text = extract_text_from_bytes(retrieved_bytes, 'xlsx')
        
        assert "Столбец A" in extracted_text
        assert "Столбец B" in extracted_text
        assert "Кириллица" in extracted_text
        assert "999" in extracted_text
        
    except ImportError:
        pytest.skip("openpyxl not available")


def test_upload_json_to_blob_and_extract(db, test_user, blob_service):
    """Тест: загрузка JSON в blob → извлечение текста из blob."""
    content = '{"название": "Тестовый JSON", "значение": 42, "список": [1, 2, 3]}'
    file_bytes = content.encode('utf-8')
    
    fake_file = FileStorage(
        stream=BytesIO(file_bytes),
        filename="test_data.json",
        content_type="application/json"
    )
    
    # Сохраняем в blob
    document, is_new = blob_service.save_file_to_db(
        db=db,
        file=fake_file,
        user_id=test_user.id,
        user_path="test/data.json"
    )
    
    assert is_new is True
    assert document.blob is not None
    
    # Извлекаем текст из blob
    retrieved_bytes = blob_service.get_file_bytes(db, document.id)
    extracted_text = extract_text_from_bytes(retrieved_bytes, 'json')
    
    assert "название" in extracted_text
    assert "Тестовый JSON" in extracted_text
    assert "значение" in extracted_text
    assert "42" in extracted_text
