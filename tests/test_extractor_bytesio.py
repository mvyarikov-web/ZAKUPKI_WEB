"""Тесты для extract_text_from_bytes - извлечение текста из бинарных данных."""
import pytest
from pathlib import Path
from document_processor.extractors.text_extractor import extract_text_from_bytes


def test_extract_txt_from_bytes_utf8():
    """Тест извлечения TXT (UTF-8) из bytes."""
    content = "Привет, мир! Это тест UTF-8 кодировки."
    file_bytes = content.encode('utf-8')
    
    result = extract_text_from_bytes(file_bytes, 'txt')
    
    assert "Привет" in result
    assert "мир" in result
    assert "UTF-8" in result


def test_extract_txt_from_bytes_cp1251():
    """Тест извлечения TXT (CP1251) из bytes."""
    content = "Тестовый документ в кодировке Windows-1251"
    file_bytes = content.encode('cp1251')
    
    result = extract_text_from_bytes(file_bytes, 'txt')
    
    assert "Тестовый" in result
    assert "документ" in result
    assert "Windows-1251" in result


def test_extract_json_from_bytes():
    """Тест извлечения JSON из bytes."""
    content = '{"key": "значение", "number": 42}'
    file_bytes = content.encode('utf-8')
    
    result = extract_text_from_bytes(file_bytes, 'json')
    
    assert "key" in result
    assert "значение" in result
    assert "42" in result


def test_extract_html_from_bytes():
    """Тест извлечения HTML из bytes с очисткой тегов."""
    content = """
    <html>
    <head><title>Тест</title></head>
    <body>
        <h1>Заголовок</h1>
        <p>Параграф с <b>жирным</b> текстом.</p>
        <script>alert('remove me');</script>
    </body>
    </html>
    """
    file_bytes = content.encode('utf-8')
    
    result = extract_text_from_bytes(file_bytes, 'html')
    
    assert "Заголовок" in result
    assert "Параграф" in result
    assert "жирным" in result
    assert "alert" not in result  # Скрипты должны быть удалены
    assert "<b>" not in result    # Теги должны быть удалены


def test_extract_pdf_from_bytes(tmp_path):
    """Тест извлечения PDF из bytes."""
    # Используем существующий тестовый PDF или создаём минимальный
    test_pdf = tmp_path / "test.pdf"
    
    # Создаём простой PDF с помощью pypdf (если доступен)
    try:
        from pypdf import PdfWriter
        from io import BytesIO
        
        writer = PdfWriter()
        # Создаём пустую страницу (минимальный PDF)
        writer.add_blank_page(width=200, height=200)
        
        # Сохраняем в bytes
        pdf_buffer = BytesIO()
        writer.write(pdf_buffer)
        file_bytes = pdf_buffer.getvalue()
        
        result = extract_text_from_bytes(file_bytes, 'pdf')
        
        # Пустая страница может вернуть пустую строку - это ОК
        assert isinstance(result, str)
        
    except ImportError:
        pytest.skip("pypdf not available for PDF creation")


def test_extract_docx_from_bytes(tmp_path):
    """Тест извлечения DOCX из bytes."""
    try:
        from docx import Document
        from io import BytesIO
        
        # Создаём простой DOCX
        doc = Document()
        doc.add_paragraph("Первый параграф с кириллицей")
        doc.add_paragraph("Второй параграф с числами: 123")
        
        # Сохраняем в bytes
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        file_bytes = docx_buffer.getvalue()
        
        result = extract_text_from_bytes(file_bytes, 'docx')
        
        assert "Первый параграф" in result
        assert "кириллицей" in result
        assert "Второй параграф" in result
        assert "123" in result
        
    except ImportError:
        pytest.skip("python-docx not available")


def test_extract_xlsx_from_bytes(tmp_path):
    """Тест извлечения XLSX из bytes."""
    try:
        from openpyxl import Workbook
        from io import BytesIO
        
        # Создаём простой XLSX
        wb = Workbook()
        ws = wb.active
        ws['A1'] = "Заголовок"
        ws['B1'] = "Значение"
        ws['A2'] = "Строка 1"
        ws['B2'] = 42
        ws['A3'] = "Строка 2"
        ws['B3'] = "Кириллица"
        
        # Сохраняем в bytes
        xlsx_buffer = BytesIO()
        wb.save(xlsx_buffer)
        file_bytes = xlsx_buffer.getvalue()
        
        result = extract_text_from_bytes(file_bytes, 'xlsx')
        
        assert "Заголовок" in result
        assert "Значение" in result
        assert "Строка 1" in result
        assert "42" in result
        assert "Кириллица" in result
        
    except ImportError:
        pytest.skip("openpyxl not available")


def test_extract_unsupported_format_from_bytes():
    """Тест для неподдерживаемого формата - должна вернуться пустая строка."""
    file_bytes = b"random binary data \x00\xff\xfe"
    
    result = extract_text_from_bytes(file_bytes, 'unknown')
    
    assert result == ''


def test_extract_corrupted_data_from_bytes():
    """Тест для повреждённых данных - должна вернуться пустая строка (graceful degrade)."""
    file_bytes = b"\x00\xff\xfe random corrupted data"
    
    # Пытаемся извлечь как PDF (заведомо некорректный формат)
    result = extract_text_from_bytes(file_bytes, 'pdf')
    
    assert result == ''  # Ошибка должна быть поглощена


def test_extract_empty_bytes():
    """Тест для пустого файла."""
    file_bytes = b""
    
    result = extract_text_from_bytes(file_bytes, 'txt')
    
    assert result == ''
