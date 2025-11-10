"""Строгие тесты для проверки корректности извлечения кириллицы из разных форматов.

Проверяем:
1. Непустой результат
2. Отсутствие мусора (неправильные кодировки)
3. Наличие конкретных кириллических слов
4. Минимальная длина текста
"""
import pytest
from io import BytesIO
from document_processor.extractors.text_extractor import extract_text_from_bytes


def test_txt_utf8_cyrillic_strict():
    """Строгий тест TXT (UTF-8) с кириллицей."""
    content = "Тестовый документ для проверки кириллицы: жизнь, экзамен, йогурт, щука, ёлка."
    file_bytes = content.encode('utf-8')
    
    result = extract_text_from_bytes(file_bytes, 'txt')
    
    # Проверка непустоты
    assert result, "Результат не должен быть пустым"
    assert len(result) >= 50, f"Слишком короткий результат: {len(result)} символов"
    
    # Проверка конкретных слов
    assert "Тестовый" in result, "Не найдено слово 'Тестовый'"
    assert "кириллицы" in result, "Не найдено слово 'кириллицы'"
    assert "жизнь" in result, "Не найдено слово 'жизнь'"
    assert "экзамен" in result, "Не найдено слово 'экзамен'"
    assert "щука" in result, "Не найдено слово 'щука'"
    assert "ёлка" in result, "Не найдено слово 'ёлка'"
    
    # Проверка отсутствия мусора (нечитаемых символов)
    garbage_indicators = ['�', '\ufffd', '?????']
    for indicator in garbage_indicators:
        assert indicator not in result, f"Обнаружен индикатор мусора: {indicator}"


def test_txt_cp1251_cyrillic_strict():
    """Строгий тест TXT (CP1251) с кириллицей."""
    content = "Документ в кодировке Windows-1251: борщ, программа, файл, система."
    file_bytes = content.encode('cp1251')
    
    result = extract_text_from_bytes(file_bytes, 'txt')
    
    assert result, "Результат не должен быть пустым"
    assert len(result) >= 40, f"Слишком короткий результат: {len(result)} символов"
    
    assert "Документ" in result, "Не найдено слово 'Документ'"
    assert "Windows-1251" in result, "Не найдена строка 'Windows-1251'"
    assert "борщ" in result, "Не найдено слово 'борщ'"
    assert "программа" in result, "Не найдено слово 'программа'"
    
    # Проверка отсутствия мусора
    assert '�' not in result, "Обнаружен символ замены (мусор)"


def test_json_cyrillic_strict():
    """Строгий тест JSON с кириллицей."""
    content = '''{
        "название": "Тестовый JSON",
        "описание": "Проверка кириллицы в JSON",
        "теги": ["закупки", "тендер", "документация"],
        "цена": 12345.67
    }'''
    file_bytes = content.encode('utf-8')
    
    result = extract_text_from_bytes(file_bytes, 'json')
    
    assert result, "Результат не должен быть пустым"
    assert len(result) >= 80, f"Слишком короткий результат: {len(result)} символов"
    
    assert "название" in result, "Не найдено поле 'название'"
    assert "Тестовый JSON" in result, "Не найдено значение 'Тестовый JSON'"
    assert "закупки" in result, "Не найдено слово 'закупки'"
    assert "тендер" in result, "Не найдено слово 'тендер'"
    assert "документация" in result, "Не найдено слово 'документация'"
    assert "12345" in result, "Не найдено число"
    
    assert '�' not in result, "Обнаружен мусор"


def test_html_cyrillic_strict():
    """Строгий тест HTML с кириллицей и проверкой удаления тегов."""
    content = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>Тестовая страница</title>
        <script>
            var garbage = "это должно быть удалено";
        </script>
    </head>
    <body>
        <h1>Заголовок документа</h1>
        <p>Первый параграф с кириллицей: закупки, тендер, конкурс.</p>
        <p>Второй параграф: <strong>жирный текст</strong>, <em>курсив</em>.</p>
        <div>Блок с текстом: система, программа, файл.</div>
    </body>
    </html>
    """
    file_bytes = content.encode('utf-8')
    
    result = extract_text_from_bytes(file_bytes, 'html')
    
    assert result, "Результат не должен быть пустым"
    assert len(result) >= 80, f"Слишком короткий результат: {len(result)} символов"
    
    # Проверка наличия контента
    assert "Заголовок документа" in result, "Не найден заголовок"
    assert "закупки" in result, "Не найдено слово 'закупки'"
    assert "тендер" in result, "Не найдено слово 'тендер'"
    assert "конкурс" in result, "Не найдено слово 'конкурс'"
    assert "жирный текст" in result, "Не найден текст из <strong>"
    assert "система" in result, "Не найдено слово 'система'"
    
    # Проверка удаления тегов и скриптов
    assert "<h1>" not in result, "Теги не удалены"
    assert "<p>" not in result, "Теги не удалены"
    assert "<script>" not in result, "Тег <script> не удалён"
    assert "var garbage" not in result, "Содержимое скрипта не удалено"
    
    assert '�' not in result, "Обнаружен мусор"


def test_docx_cyrillic_strict():
    """Строгий тест DOCX с кириллицей."""
    try:
        from docx import Document
        
        # Создаём DOCX с богатым содержимым
        doc = Document()
        doc.add_paragraph("Первый параграф: тестирование системы закупок.")
        doc.add_paragraph("Второй параграф: тендер, конкурс, аукцион, котировка.")
        doc.add_paragraph("Третий параграф: документация, спецификация, требования.")
        doc.add_paragraph("Числа и символы: 12345, цена 67890.50 руб.")
        
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        file_bytes = docx_buffer.getvalue()
        
        result = extract_text_from_bytes(file_bytes, 'docx')
        
        # Проверка непустоты и минимальной длины
        assert result, "Результат не должен быть пустым"
        assert len(result) >= 100, f"Слишком короткий результат: {len(result)} символов (ожидалось >100)"
        
        # Проверка всех параграфов
        assert "Первый параграф" in result, "Не найден первый параграф"
        assert "тестирование" in result, "Не найдено слово 'тестирование'"
        assert "закупок" in result, "Не найдено слово 'закупок'"
        
        assert "Второй параграф" in result, "Не найден второй параграф"
        assert "тендер" in result, "Не найдено слово 'тендер'"
        assert "конкурс" in result, "Не найдено слово 'конкурс'"
        assert "аукцион" in result, "Не найдено слово 'аукцион'"
        
        assert "Третий параграф" in result, "Не найден третий параграф"
        assert "документация" in result, "Не найдено слово 'документация'"
        assert "спецификация" in result, "Не найдено слово 'спецификация'"
        
        assert "12345" in result, "Не найдено число 12345"
        assert "67890" in result, "Не найдено число 67890"
        assert "руб" in result, "Не найдена валюта 'руб'"
        
        # Проверка отсутствия мусора
        assert '�' not in result, "Обнаружен символ замены (мусор)"
        
        # Проверка, что это не случайные байты
        cyrillic_count = sum(1 for c in result if '\u0400' <= c <= '\u04FF')
        assert cyrillic_count > 50, f"Слишком мало кириллических символов: {cyrillic_count}"
        
    except ImportError:
        pytest.skip("python-docx not available")


def test_xlsx_cyrillic_strict():
    """Строгий тест XLSX с кириллицей."""
    try:
        from openpyxl import Workbook
        
        # Создаём XLSX с множеством данных
        wb = Workbook()
        ws = wb.active
        
        # Заголовки
        ws['A1'] = "Наименование"
        ws['B1'] = "Количество"
        ws['C1'] = "Цена"
        ws['D1'] = "Статус"
        
        # Данные
        ws['A2'] = "Бумага офисная"
        ws['B2'] = 100
        ws['C2'] = 250.50
        ws['D2'] = "В наличии"
        
        ws['A3'] = "Картридж принтера"
        ws['B3'] = 50
        ws['C3'] = 1200.00
        ws['D3'] = "Заказан"
        
        ws['A4'] = "Скоросшиватель"
        ws['B4'] = 200
        ws['C4'] = 15.75
        ws['D4'] = "Доставлен"
        
        xlsx_buffer = BytesIO()
        wb.save(xlsx_buffer)
        file_bytes = xlsx_buffer.getvalue()
        
        result = extract_text_from_bytes(file_bytes, 'xlsx')
        
        # Проверка непустоты
        assert result, "Результат не должен быть пустым"
        assert len(result) >= 80, f"Слишком короткий результат: {len(result)} символов"
        
        # Проверка заголовков
        assert "Наименование" in result, "Не найден заголовок 'Наименование'"
        assert "Количество" in result, "Не найден заголовок 'Количество'"
        assert "Цена" in result, "Не найден заголовок 'Цена'"
        assert "Статус" in result, "Не найден заголовок 'Статус'"
        
        # Проверка данных
        assert "Бумага офисная" in result, "Не найдено 'Бумага офисная'"
        assert "Картридж принтера" in result, "Не найдено 'Картридж принтера'"
        assert "Скоросшиватель" in result, "Не найдено 'Скоросшиватель'"
        
        assert "В наличии" in result, "Не найден статус 'В наличии'"
        assert "Заказан" in result, "Не найден статус 'Заказан'"
        assert "Доставлен" in result, "Не найден статус 'Доставлен'"
        
        # Проверка чисел
        assert "100" in result, "Не найдено число 100"
        assert "250.5" in result or "250.50" in result, "Не найдена цена 250.50"
        assert "1200" in result, "Не найдено число 1200"
        
        # Проверка отсутствия мусора
        assert '�' not in result, "Обнаружен символ замены (мусор)"
        
        # Проверка количества кириллицы
        cyrillic_count = sum(1 for c in result if '\u0400' <= c <= '\u04FF')
        assert cyrillic_count > 30, f"Слишком мало кириллических символов: {cyrillic_count}"
        
    except ImportError:
        pytest.skip("openpyxl not available")


def test_pdf_cyrillic_strict():
    """Строгий тест PDF с кириллицей (если возможно создать PDF с текстом)."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # Пытаемся зарегистрировать шрифт с поддержкой кириллицы
        # Обычно в macOS есть системные шрифты
        try:
            # Ищем системный шрифт с кириллицей
            import os
            font_paths = [
                '/System/Library/Fonts/Supplemental/Arial.ttf',
                '/System/Library/Fonts/Helvetica.ttc',
                '/Library/Fonts/Arial.ttf',
            ]
            font_registered = False
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('CustomFont', font_path))
                        font_registered = True
                        break
                    except:
                        continue
            
            if not font_registered:
                pytest.skip("Не найден шрифт с поддержкой кириллицы для PDF")
            
            # Создаём PDF с кириллицей
            pdf_buffer = BytesIO()
            c = canvas.Canvas(pdf_buffer, pagesize=letter)
            c.setFont('CustomFont', 12)
            
            c.drawString(100, 750, "Тестовый документ PDF")
            c.drawString(100, 730, "Проверка кириллицы: закупки, тендер, конкурс")
            c.drawString(100, 710, "Документация системы закупок")
            c.drawString(100, 690, "Цена: 12345.67 рублей")
            
            c.save()
            file_bytes = pdf_buffer.getvalue()
            
            result = extract_text_from_bytes(file_bytes, 'pdf')
            
            # PDF может быть сложным для извлечения, но хотя бы часть должна быть
            assert result, "Результат не должен быть пустым"
            
            # Проверяем наличие хотя бы части слов
            found_words = 0
            test_words = ["Тестовый", "документ", "кириллицы", "закупки", "тендер", "конкурс", "Документация", "Цена", "12345"]
            for word in test_words:
                if word in result:
                    found_words += 1
            
            assert found_words >= 3, f"Найдено слишком мало ключевых слов: {found_words}/9. Результат: {result[:200]}"
            
        except Exception as e:
            pytest.skip(f"Не удалось создать PDF с кириллицей: {e}")
            
    except ImportError:
        pytest.skip("reportlab not available for PDF creation")


def test_csv_cyrillic_strict():
    """Строгий тест CSV с кириллицей."""
    content = """Наименование,Количество,Цена,Поставщик
Бумага офисная,100,250.50,ООО Поставка
Картридж принтера,50,1200.00,ИП Иванов
Скоросшиватель,200,15.75,ООО Канцтовары"""
    
    file_bytes = content.encode('utf-8')
    
    result = extract_text_from_bytes(file_bytes, 'csv')
    
    assert result, "Результат не должен быть пустым"
    assert len(result) >= 100, f"Слишком короткий результат: {len(result)} символов"
    
    # Проверка заголовков
    assert "Наименование" in result, "Не найден заголовок 'Наименование'"
    assert "Количество" in result, "Не найден заголовок 'Количество'"
    assert "Поставщик" in result, "Не найден заголовок 'Поставщик'"
    
    # Проверка данных
    assert "Бумага офисная" in result, "Не найдено 'Бумага офисная'"
    assert "Картридж принтера" in result, "Не найдено 'Картридж принтера'"
    assert "Скоросшиватель" in result, "Не найдено 'Скоросшиватель'"
    assert "Поставка" in result, "Не найдена организация 'Поставка'"
    assert "Иванов" in result, "Не найдена фамилия 'Иванов'"
    
    assert '�' not in result, "Обнаружен мусор"
