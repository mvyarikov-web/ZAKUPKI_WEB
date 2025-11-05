"""Тесты для групповой индексации (increment-014)."""
import os
import pytest
from document_processor.search.indexer import Indexer
from document_processor.search.searcher import Searcher


def test_grouped_indexing_creates_skeleton(tmp_path):
    """Проверяем, что групповая индексация создаёт скелет с маркерами групп."""
    # Создаём тестовые файлы разных типов
    test_dir = tmp_path / "test_docs"
    test_dir.mkdir()
    
    # Быстрый файл
    (test_dir / "fast.txt").write_text("быстрый файл с текстом", encoding='utf-8')
    
    # Средний файл (DOCX)
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("средний файл")
        docx_path = test_dir / "medium.docx"
        doc.save(str(docx_path))
    except ImportError:
        pytest.skip("python-docx not available")
    
    # Создаём индекс с группировкой
    indexer = Indexer()
    index_path = indexer.create_index(str(test_dir), use_groups=True)
    
    assert os.path.exists(index_path)
    
    # Проверяем содержимое индекса
    with open(index_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Должны быть заголовки групп
    assert '[ГРУППА: FAST]' in content
    assert '[ГРУППА: MEDIUM]' in content
    assert '[ГРУППА: SLOW]' in content
    
    # Должны быть маркеры BEGIN/END
    assert '<!-- BEGIN_FAST -->' in content
    assert '<!-- END_FAST -->' in content
    assert '<!-- BEGIN_MEDIUM -->' in content
    assert '<!-- END_MEDIUM -->' in content
    
    # Статусы должны быть "завершено"
    assert '✅ завершено' in content


def test_grouped_indexing_content_in_sections(tmp_path):
    """Проверяем, что содержимое файлов попадает в соответствующие секции."""
    test_dir = tmp_path / "test_docs"
    test_dir.mkdir()
    
    # Быстрый файл
    (test_dir / "fast.txt").write_text("быстрое содержимое жираф", encoding='utf-8')
    
    # Средний файл (DOCX) - создаём через библиотеку
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("содержимое со словом слон")
        docx_path = test_dir / "medium.docx"
        doc.save(str(docx_path))
    except ImportError:
        pytest.skip("python-docx not available")
    
    # Создаём индекс с группировкой
    indexer = Indexer()
    index_path = indexer.create_index(str(test_dir), use_groups=True)
    
    with open(index_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Проверяем, что термины присутствуют
    assert 'жираф' in content
    assert 'слон' in content
    
    # Проверяем структуру: содержимое между маркерами
    import re
    
    # Извлекаем секцию FAST
    fast_match = re.search(r'<!-- BEGIN_FAST -->(.*?)<!-- END_FAST -->', content, re.DOTALL)
    assert fast_match is not None
    fast_section = fast_match.group(1)
    assert 'жираф' in fast_section
    assert 'fast.txt' in fast_section
    
    # Извлекаем секцию MEDIUM
    medium_match = re.search(r'<!-- BEGIN_MEDIUM -->(.*?)<!-- END_MEDIUM -->', content, re.DOTALL)
    assert medium_match is not None
    medium_section = medium_match.group(1)
    assert 'слон' in medium_section
    assert 'medium.docx' in medium_section


def test_searcher_ignores_group_headers(tmp_path):
    """Проверяем, что Searcher игнорирует заголовки групп при поиске."""
    test_dir = tmp_path / "test_docs"
    test_dir.mkdir()
    
    # Создаём файл со словом "ГРУППА" в содержимом
    (test_dir / "test.txt").write_text("документ содержит слово ГРУППА и другой текст", encoding='utf-8')
    
    # Создаём индекс с группировкой
    indexer = Indexer()
    index_path = indexer.create_index(str(test_dir), use_groups=True)
    
    # Ищем слово "ГРУППА"
    searcher = Searcher()
    results = searcher.search(index_path, ['ГРУППА'], context=50)
    
    # Должно быть найдено только одно вхождение (из документа, не из заголовков)
    assert len(results) == 1
    assert 'test.txt' in results[0]['title']
    assert 'другой текст' in results[0]['snippet']


def test_classic_indexing_still_works(tmp_path):
    """Проверяем, что классическая индексация (без групп) продолжает работать."""
    test_dir = tmp_path / "test_docs"
    test_dir.mkdir()
    
    (test_dir / "doc1.txt").write_text("первый документ", encoding='utf-8')
    (test_dir / "doc2.txt").write_text("второй документ", encoding='utf-8')
    
    # Создаём индекс без группировки (классический режим)
    indexer = Indexer()
    index_path = indexer.create_index(str(test_dir), use_groups=False)
    
    assert os.path.exists(index_path)
    
    with open(index_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # НЕ должно быть заголовков групп
    assert '[ГРУППА:' not in content
    assert '<!-- BEGIN_' not in content
    
    # Должно быть содержимое
    assert 'первый документ' in content
    assert 'второй документ' in content
    
    # Поиск должен работать
    searcher = Searcher()
    results = searcher.search(index_path, ['первый'], context=50)
    assert len(results) == 1
    assert 'doc1.txt' in results[0]['title']


def test_file_classification(tmp_path):
    """Проверяем правильность классификации файлов по группам."""
    test_dir = tmp_path / "test_docs"
    test_dir.mkdir()
    
    # Создаём файлы разных типов
    (test_dir / "fast.txt").write_text("текст", encoding='utf-8')
    (test_dir / "fast.csv").write_text("a,b", encoding='utf-8')
    (test_dir / "fast.html").write_text("<html>тест</html>", encoding='utf-8')
    
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("docx content")
        doc.save(str(test_dir / "medium.docx"))
    except ImportError:
        # Если библиотеки нет, создадим простой ZIP
        import zipfile
        docx_path = test_dir / "medium.docx"
        with zipfile.ZipFile(docx_path, 'w') as z:
            z.writestr("word/document.xml", 'dummy')
    
    # Создадим простой ZIP архив
    import zipfile
    with zipfile.ZipFile(test_dir / "slow.zip", 'w') as z:
        z.writestr("test.txt", "content")
    
    indexer = Indexer()
    all_files = indexer._collect_all_files(str(test_dir))
    groups = indexer._classify_files(all_files)
    
    # Проверяем количество файлов в каждой группе
    assert len(groups['fast']) == 3  # txt, csv, html
    assert len(groups['medium']) == 1  # docx
    assert len(groups['slow']) == 1  # zip


def test_grouped_index_status_tracking(tmp_path):
    """Проверяем, что статус индексации отслеживает группы."""
    test_dir = tmp_path / "test_docs"
    test_dir.mkdir()
    
    # Создаём структуру для статуса
    index_folder = tmp_path / "index"
    index_folder.mkdir()
    index_folder / "status.json"
    
    # Создаём тестовые файлы
    (test_dir / "fast.txt").write_text("текст", encoding='utf-8')
    
    # Создаём индекс с группировкой
    # Нужно переместить test_dir внутрь папки с правильной структурой
    # чтобы статус записался в правильное место
    parent = tmp_path / "root"
    parent.mkdir()
    uploads = parent / "uploads"
    test_dir.rename(uploads)
    
    idx_folder = parent / "index"
    idx_folder.mkdir()
    
    indexer = Indexer()
    indexer.create_index(str(uploads), use_groups=True)
    
    # Проверяем, что статус был создан
    status_file = idx_folder / "status.json"
    assert status_file.exists()
    
    import json
    with open(status_file, 'r', encoding='utf-8') as f:
        status = json.load(f)
    
    # Проверяем поля статуса
    assert status['status'] == 'completed'
    assert 'group_status' in status
    assert status['group_status']['fast'] == 'completed'


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
