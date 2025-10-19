from __future__ import annotations
import os
import io
import re
import zipfile
import tempfile
import datetime
import json
from datetime import datetime
from typing import Iterable, Tuple, List, Optional, Any
import logging

# Разделитель заголовков в индексе
HEADER_BAR = "===============================" 
# Маркеры начала и конца документа
DOC_START_MARKER = "<<< НАЧАЛО ДОКУМЕНТА >>>"
DOC_END_MARKER = "<<< КОНЕЦ ДОКУМЕНТА >>>"
SUPPORTED_EXT = {"pdf", "doc", "docx", "xls", "xlsx", "txt", "zip", "rar", "html", "htm", "csv", "tsv", "xml", "json"}
DOC_EXTS = {"doc", "docx"}

class Indexer:
    def __init__(self, *, max_depth: int = 10, archive_depth: int = 0):
        self.max_depth = max_depth
        self.archive_depth = archive_depth
        self._temp_paths: List[str] = []
        self._log = logging.getLogger(__name__)

    def _classify_files(self, files: List[Tuple[str, str, str, str]]) -> dict:
        """Группирует файлы по скорости обработки.
        
        Args:
            files: список кортежей (ext, name, abs_path, rel_path)
        
        Returns:
            {
                'fast': [(rel_path, abs_path), ...],    # TXT, CSV, HTML
                'medium': [...],                         # DOCX, XLSX, PDF (попытка без OCR)
                'slow': [...]                            # PDF с OCR, архивы
            }
        """
        fast_exts = {'.txt', '.csv', '.html', '.htm', '.md', '.json', '.xml', '.tsv'}
        medium_exts = {'.docx', '.xlsx', '.xls', '.doc'}
        
        groups = {'fast': [], 'medium': [], 'slow': []}
        
        for ext, name, abs_path, rel_path in files:
            ext_lower = f'.{ext}'
            
            if ext_lower in fast_exts:
                groups['fast'].append((rel_path, abs_path))
            elif ext_lower in medium_exts:
                groups['medium'].append((rel_path, abs_path))
            elif ext == 'pdf':
                # Эвристика: пробуем определить, текстовый ли PDF
                if self._is_text_pdf(abs_path):
                    groups['medium'].append((rel_path, abs_path))
                else:
                    groups['slow'].append((rel_path, abs_path))
            else:  # ZIP, RAR и прочие
                groups['slow'].append((rel_path, abs_path))
        
        return groups
    
    def _is_text_pdf(self, path: str, threshold: int = 100) -> bool:
        """Быстрая проверка: содержит ли PDF извлекаемый текст.
        
        Args:
            path: Путь к PDF
            threshold: Минимум символов для классификации как "текстовый"
        
        Returns:
            True если PDF содержит текст ≥ threshold символов
        """
        try:
            # Пробуем извлечь первую страницу через pdfplumber (быстро)
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                if pdf.pages:
                    text = pdf.pages[0].extract_text() or ''
                    return len(text.strip()) >= threshold
        except Exception:
            pass
        
        return False  # По умолчанию считаем, что нужен OCR

    def create_index(self, root_folder: str, use_groups: bool = False) -> str:
        """Создаёт индекс для файлов в указанной папке.
        
        Args:
            root_folder: корневая папка для индексации
            use_groups: если True, использует групповую индексацию (increment-014)
        
        Returns:
            путь к созданному индексу
        """
        if use_groups:
            return self._create_index_grouped(root_folder)
        else:
            return self._create_index_classic(root_folder)
    
    def _create_index_grouped(self, root_folder: str) -> str:
        """Создаёт индекс в 3 этапа с промежуточной доступностью (increment-014)."""
        index_path = os.path.join(root_folder, "_search_index.txt")
        status_path = os.path.join(os.path.dirname(root_folder), "index", "status.json")
        
        self._log.info("Групповая индексация начата: root=%s -> %s", root_folder, index_path)
        
        # 1. Собираем и классифицируем файлы
        all_files = self._collect_all_files(root_folder)
        groups = self._classify_files(all_files)
        total_files = len(all_files)
        
        # Создаём директорию для статуса, если не существует
        os.makedirs(os.path.dirname(status_path), exist_ok=True)
        
        try:
            # Инициализация статуса
            self._update_status(status_path, {
                'status': 'running',
                'total': total_files,
                'processed': 0,
                'current_group': None,
                'group_status': {
                    'fast': 'pending',
                    'medium': 'pending',
                    'slow': 'pending'
                },
                'started_at': datetime.now().isoformat(),
                'updated_at': datetime.now().isoformat()
            })
            
            # 2. Создаём индекс с заголовками
            self._create_index_skeleton(index_path, groups)
            
            # 3. Обрабатываем группы последовательно
            processed_files = 0
            for group_name in ['fast', 'medium', 'slow']:
                group_files = groups[group_name]
                if not group_files:
                    continue
                
                self._log.info(f"Обработка группы {group_name}: {len(group_files)} файлов")
                
                # Обновляем статус группы
                self._update_status(status_path, {
                    'status': 'running',
                    'total': total_files,
                    'processed': processed_files,
                    'current_group': group_name,
                    'group_status': {
                        'fast': 'completed' if group_name != 'fast' else 'running',
                        'medium': 'completed' if group_name == 'slow' else ('running' if group_name == 'medium' else 'pending'),
                        'slow': 'running' if group_name == 'slow' else 'pending'
                    },
                    'updated_at': datetime.now().isoformat()
                })
                
                # Обрабатываем группу во временный файл
                temp_file = os.path.join(os.path.dirname(status_path), f'_search_index_{group_name}.tmp')
                self._process_group(group_files, temp_file, group_name)
                
                # Вставляем в главный индекс атомарно
                self._insert_group_into_index(index_path, group_name, temp_file)
                
                # Удаляем временный файл
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                
                processed_files += len(group_files)
                
                # Обновляем статус после завершения группы
                self._update_status(status_path, {
                    'processed': processed_files,
                    'group_status': {
                        'fast': 'completed',
                        'medium': 'completed' if group_name != 'fast' else 'pending',
                        'slow': 'completed' if group_name == 'slow' else 'pending'
                    },
                    'updated_at': datetime.now().isoformat()
                })
            
            # 4. Финализация
            self._update_status(status_path, {
                'status': 'completed',
                'total': total_files,
                'processed': total_files,
                'current_group': None,
                'completed_at': datetime.now().isoformat(),
                'updated_at': datetime.now().isoformat()
            })
            
            self._log.info("Групповая индексация завершена: %s", index_path)
            return index_path
        
        except Exception as e:
            # Обновляем статус об ошибке
            self._update_status(status_path, {
                'status': 'error',
                'error': str(e),
                'updated_at': datetime.now().isoformat()
            })
            
            self._log.exception("Ошибка при групповой индексации: %s", e)
            raise
        finally:
            self._cleanup_temp_paths()
    
    def _create_index_classic(self, root_folder: str) -> str:
        """Классическая индексация (одним проходом) — для обратной совместимости."""
        index_path = os.path.join(root_folder, "_search_index.txt")
        temp_path = os.path.join(root_folder, "._search_index.tmp")
        status_path = os.path.join(os.path.dirname(root_folder), "index", "status.json")
        
        self._log.info("Индексация начата: root=%s -> %s", root_folder, index_path)
        
        # Подсчитываем общее количество файлов для прогресса
        all_files = list(self._collect_all_files(root_folder))
        total_files = len(all_files)
        processed_files = 0
        
        # Создаём директорию для статуса, если не существует
        os.makedirs(os.path.dirname(status_path), exist_ok=True)
        
        try:
            # Инициализация статуса
            self._update_status(status_path, {
                'status': 'running',
                'total': total_files,
                'processed': 0,
                'current_file': None,
                'current_format': None,
                'ocr_active': False,
                'started_at': datetime.now().isoformat(),
                'updated_at': datetime.now().isoformat()
            })
            
            # Write to temporary file first for atomic replacement
            with io.open(temp_path, "w", encoding="utf-8") as out:
                for rel_path, abs_path, source in self._iter_sources(root_folder):
                    # Обновляем статус для текущего файла
                    ext = rel_path.rsplit(".", 1)[-1].upper() if "." in rel_path else "UNKNOWN"
                    is_pdf = ext == "PDF"
                    
                    self._update_status(status_path, {
                        'status': 'running',
                        'total': total_files,
                        'processed': processed_files,
                        'current_file': os.path.basename(rel_path),
                        'current_format': ext,
                        'ocr_active': is_pdf,  # PDF может запустить OCR
                        'updated_at': datetime.now().isoformat()
                    })
                    
                    text, meta = self._extract_text(abs_path, rel_path, source)
                    self._write_entry(out, rel_path=source, text=text, meta=meta)
                    processed_files += 1
            
            # Atomically replace old index with new one
            if os.path.exists(temp_path):
                os.replace(temp_path, index_path)
                self._log.info("Индексация завершена: %s", index_path)
            else:
                self._log.error("Временный файл индекса не создан: %s", temp_path)
                raise RuntimeError(f"Failed to create temporary index file: {temp_path}")
            
            # Финальное обновление статуса
            self._update_status(status_path, {
                'status': 'completed',
                'total': total_files,
                'processed': total_files,
                'current_file': None,
                'current_format': None,
                'ocr_active': False,
                'completed_at': datetime.now().isoformat(),
                'updated_at': datetime.now().isoformat()
            })
            
            return index_path
        except Exception as e:
            # Clean up temporary file on error
            try:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            except Exception:
                pass
            
            # Обновляем статус об ошибке
            self._update_status(status_path, {
                'status': 'error',
                'error': str(e),
                'updated_at': datetime.now().isoformat()
            })
            
            self._log.exception("Ошибка при создании индекса: %s", e)
            raise
        finally:
            self._cleanup_temp_paths()

    def _iter_sources(self, root_folder: str) -> Iterable[Tuple[str, str, str]]:
        """Iterate over sources, sorted by processing complexity (fast → slow).
        
        Processing order:
        1. TXT, CSV, HTML (instant)
        2. DOCX, XLSX, vector PDFs (fast)
        3. PDF scans with OCR, ZIP, RAR (slow)
        """
        # Collect all files first for sorting
        all_files = []
        
        # Walk directories recursively
        for dirpath, dirnames, filenames in os.walk(root_folder):
            # Optionally limit depth
            rel_dir = os.path.relpath(dirpath, root_folder)
            depth = 0 if rel_dir == "." else rel_dir.count(os.sep) + 1
            if depth > self.max_depth:
                continue
            for name in filenames:
                # Пропускаем временные файлы Office (обычно начинаются с ~$ или $)
                if name.startswith("~$") or name.startswith("$"):
                    continue
                # Пропускаем сам сводный индекс
                if name == "_search_index.txt":
                    continue
                ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
                if ext in SUPPORTED_EXT:
                    abs_path = os.path.join(dirpath, name)
                    rel_path = os.path.normpath(os.path.join(rel_dir, name)) if rel_dir != "." else name
                    all_files.append((ext, name, abs_path, rel_path))
        
        # Sort files by processing priority
        all_files.sort(key=self._file_sort_key)
        
        # Yield files in sorted order
        for ext, name, abs_path, rel_path in all_files:
            if ext in {"zip", "rar"}:
                for v_rel, v_abs, v_source in self._iter_archive(abs_path, rel_path, ext, current_depth=0):
                    yield v_rel, v_abs, v_source
            else:
                yield rel_path, abs_path, rel_path
    
    def _file_sort_key(self, file_info: Tuple[str, str, str, str]) -> Tuple[int, int, str]:
        """Compute sort key for file: (priority, size_category, name).
        
        Priority groups:
        - 0: instant (txt, csv, html, xml, json)
        - 1: fast (docx, xlsx)
        - 2: medium (pdf - may be vector or scan)
        - 3: slow (archives: zip, rar)
        
        Size categories: 0=small (<1MB), 1=medium (1-10MB), 2=large (>10MB)
        """
        ext, name, abs_path, rel_path = file_info
        
        # Priority by extension
        if ext in {"txt", "csv", "tsv", "html", "htm", "xml", "json"}:
            priority = 0  # instant
        elif ext in {"docx", "xlsx", "doc", "xls"}:
            priority = 1  # fast
        elif ext == "pdf":
            priority = 2  # medium (could be vector or OCR)
        else:  # zip, rar
            priority = 3  # slow
        
        # Size category (to process small files first within same priority)
        try:
            size = os.path.getsize(abs_path)
            if size < 1_000_000:  # <1MB
                size_cat = 0
            elif size < 10_000_000:  # 1-10MB
                size_cat = 1
            else:  # >10MB
                size_cat = 2
        except Exception:
            size_cat = 0  # default to small if can't get size
        
        return (priority, size_cat, name.lower())
    
    def _collect_all_files(self, root_folder: str) -> List[Tuple[str, str, str, str]]:
        """Collect all files for counting (for progress tracking).
        
        Returns list of (ext, name, abs_path, rel_path) tuples.
        """
        all_files = []
        
        for dirpath, dirnames, filenames in os.walk(root_folder):
            rel_dir = os.path.relpath(dirpath, root_folder)
            depth = 0 if rel_dir == "." else rel_dir.count(os.sep) + 1
            if depth > self.max_depth:
                continue
            
            for name in filenames:
                if name.startswith("~$") or name.startswith("$"):
                    continue
                if name == "_search_index.txt":
                    continue
                    
                ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
                if ext in SUPPORTED_EXT:
                    abs_path = os.path.join(dirpath, name)
                    rel_path = os.path.normpath(os.path.join(rel_dir, name)) if rel_dir != "." else name
                    all_files.append((ext, name, abs_path, rel_path))
        
        return all_files
    
    def _create_index_skeleton(self, index_path: str, groups: dict) -> None:
        """Создаёт индекс с заголовками групп (резервация мест).
        
        Args:
            index_path: Путь к индексному файлу
            groups: Словарь групп {'fast': [...], 'medium': [...], 'slow': [...]}
        """
        group_labels = {
            'fast': 'TXT, CSV, HTML',
            'medium': 'DOCX, XLSX, векторные PDF',
            'slow': 'PDF-сканы с OCR'
        }
        
        temp_path = index_path + '.tmp'
        
        with io.open(temp_path, 'w', encoding='utf-8') as f:
            for group_name in ['fast', 'medium', 'slow']:
                f.write('\n')
                f.write('═' * 80 + '\n')
                f.write(f'[ГРУППА: {group_name.upper()}] {group_labels[group_name]}\n')
                f.write(f'Файлов: {len(groups[group_name])} | Статус: ожидание\n')
                f.write('═' * 80 + '\n')
                f.write(f'<!-- BEGIN_{group_name.upper()} -->\n')
                f.write(f'<!-- END_{group_name.upper()} -->\n')
                f.write('\n')
        
        os.replace(temp_path, index_path)
    
    def _process_group(self, files: List[Tuple[str, str]], temp_file: str, group_name: str) -> None:
        """Обрабатывает группу файлов и записывает во временный файл.
        
        Args:
            files: список кортежей (rel_path, abs_path)
            temp_file: путь к временному файлу для записи
            group_name: имя группы (fast/medium/slow)
        """
        os.makedirs(os.path.dirname(temp_file), exist_ok=True)
        
        with io.open(temp_file, 'w', encoding='utf-8') as out:
            for rel_path, abs_path in files:
                text, meta = self._extract_text(abs_path, rel_path, rel_path)
                self._write_entry(out, rel_path=rel_path, text=text, meta=meta)
    
    def _insert_group_into_index(self, index_path: str, group_name: str, temp_file: str) -> None:
        """Атомарно вставляет содержимое группы в главный индекс.
        
        Args:
            index_path: путь к главному индексу
            group_name: имя группы (fast/medium/slow)
            temp_file: путь к временному файлу с содержимым группы
        """
        # Читаем обработанные записи группы
        with io.open(temp_file, 'r', encoding='utf-8') as f:
            group_content = f.read()
        
        # Читаем главный индекс
        with io.open(index_path, 'r', encoding='utf-8') as f:
            index_content = f.read()
        
        # Заменяем секцию группы
        begin_marker = f'<!-- BEGIN_{group_name.upper()} -->'
        end_marker = f'<!-- END_{group_name.upper()} -->'
        
        pattern = re.compile(
            re.escape(begin_marker) + r'.*?' + re.escape(end_marker),
            re.DOTALL
        )
        
        new_content = pattern.sub(
            f'{begin_marker}\n{group_content}\n{end_marker}',
            index_content
        )
        
        # Обновляем статус группы в заголовке
        new_content = re.sub(
            rf'(\[ГРУППА: {group_name.upper()}\].*?Статус: )ожидание',
            r'\1✅ завершено',
            new_content,
            flags=re.DOTALL
        )
        
        # Атомарная запись
        temp_index = index_path + '.tmp'
        with io.open(temp_index, 'w', encoding='utf-8') as f:
            f.write(new_content)
        
        os.replace(temp_index, index_path)

    def _update_status(self, status_path: str, status_data: dict) -> None:
        """Update indexing status JSON file.
        
        Args:
            status_path: Path to status.json
            status_data: Status data to merge with existing status
        """
        try:
            # Read existing status if present
            existing = {}
            if os.path.exists(status_path):
                try:
                    with open(status_path, 'r', encoding='utf-8') as f:
                        existing = json.load(f)
                except Exception:
                    pass  # ignore errors reading old status
            
            # Merge with new data
            existing.update(status_data)
            
            # Write atomically
            temp_status = status_path + '.tmp'
            with open(temp_status, 'w', encoding='utf-8') as f:
                json.dump(existing, f, ensure_ascii=False, indent=2)
            
            os.replace(temp_status, status_path)
        except Exception as e:
            self._log.debug("Не удалось обновить статус индексации: %s", e)

    def _iter_archive(self, archive_path: str, rel_path: str, kind: str, current_depth: int = 0) -> Iterable[Tuple[str, str, str]]:
        # Extract supported files from archive into temp files and yield paths
        # FR-008: Поддержка вложенных архивов с учетом archive_depth
        entries: List[Tuple[str, str, str]] = []
        
        # Проверяем глубину вложенности архивов
        if current_depth > self.archive_depth:
            self._log.info("Пропуск вложенного архива (превышена глубина %d): %s", self.archive_depth, rel_path)
            return []
        
        try:
            if kind == "zip":
                try:
                    with zipfile.ZipFile(archive_path) as z:
                        for zi in z.infolist():
                            if zi.is_dir():
                                continue
                            inner_name = zi.filename
                            ext = inner_name.rsplit(".", 1)[-1].lower() if "." in inner_name else ""
                            if ext and ext in SUPPORTED_EXT:
                                # FR-008: Обработка вложенных архивов
                                if ext in {"zip", "rar"} and current_depth < self.archive_depth:
                                    data = z.read(zi)
                                    tmp_path = self._write_temp_file(data, suffix=f".{ext}")
                                    nested_rel = f"{rel_path}!/{inner_name}"
                                    # Рекурсивно обрабатываем вложенный архив
                                    for nested_v_rel, nested_v_abs, nested_v_source in self._iter_archive(
                                        tmp_path, nested_rel, ext, current_depth + 1
                                    ):
                                        entries.append((nested_v_rel, nested_v_abs, nested_v_source))
                                elif ext not in {"zip", "rar"}:
                                    # Обычный файл внутри архива
                                    data = z.read(zi)
                                    tmp_path = self._write_temp_file(data, suffix=f".{ext}")
                                    v_rel = f"{rel_path}!/{inner_name}"
                                    v_source = f"zip://{rel_path}!/{inner_name}"
                                    entries.append((v_rel, tmp_path, v_source))
                except zipfile.BadZipFile:
                    self._log.warning("Повреждённый ZIP архив: %s", rel_path)
                    return []
            elif kind == "rar":
                try:
                    import rarfile  # type: ignore
                    with rarfile.RarFile(archive_path) as rf:
                        for ri in rf.infolist():
                            if ri.isdir():
                                continue
                            inner_name = ri.filename
                            ext = inner_name.rsplit(".", 1)[-1].lower() if "." in inner_name else ""
                            if ext and ext in SUPPORTED_EXT:
                                # FR-008: Обработка вложенных архивов
                                if ext in {"zip", "rar"} and current_depth < self.archive_depth:
                                    with rf.open(ri) as rfp:
                                        data = rfp.read()
                                    tmp_path = self._write_temp_file(data, suffix=f".{ext}")
                                    nested_rel = f"{rel_path}!/{inner_name}"
                                    # Рекурсивно обрабатываем вложенный архив
                                    for nested_v_rel, nested_v_abs, nested_v_source in self._iter_archive(
                                        tmp_path, nested_rel, ext, current_depth + 1
                                    ):
                                        entries.append((nested_v_rel, nested_v_abs, nested_v_source))
                                elif ext not in {"zip", "rar"}:
                                    # Обычный файл внутри архива
                                    with rf.open(ri) as rfp:
                                        data = rfp.read()
                                    tmp_path = self._write_temp_file(data, suffix=f".{ext}")
                                    v_rel = f"{rel_path}!/{inner_name}"
                                    v_source = f"rar://{rel_path}!/{inner_name}"
                                    entries.append((v_rel, tmp_path, v_source))
                except Exception:
                    # rar unsupported on system or corrupted: skip gracefully
                    self._log.warning("RAR не поддержан системой или ошибка чтения: %s", rel_path)
                    return []
        except Exception:
            self._log.exception("Ошибка при чтении архива: %s", rel_path)
            return []
        
        # FR-007: Логирование событий индексации архивов
        if entries:
            self._log.info("Обработан архив %s: извлечено %d файлов (глубина %d)", rel_path, len(entries), current_depth)
        
        return entries

    def _write_temp_file(self, data: bytes, suffix: str = "") -> str:
        import tempfile
        fd, tmp_path = tempfile.mkstemp(suffix=suffix)
        with os.fdopen(fd, "wb") as f:
            f.write(data)
        self._temp_paths.append(tmp_path)
        return tmp_path

    def _cleanup_temp_paths(self):
        for p in self._temp_paths:
            try:
                if os.path.exists(p):
                    os.remove(p)
            except Exception:
                pass
        self._temp_paths.clear()

    def _extract_text(self, abs_path: str, rel_path: str, source: str):
        ext = rel_path.rsplit(".", 1)[-1].lower() if "." in rel_path else ""
        text = ""
        ocr_used = False
        try:
            if ext == "txt":
                text = self._read_text_with_encoding(abs_path)
            elif ext in {"html", "htm"}:
                raw = self._read_text_with_encoding(abs_path)
                text = self._html_to_text(raw)
            elif ext in {"xml", "json"}:
                text = self._read_text_with_encoding(abs_path)
            elif ext in {"csv", "tsv"}:
                sep = "," if ext == "csv" else "\t"
                try:
                    import csv
                    with open(abs_path, "r", encoding="utf-8", errors="ignore") as f:
                        rows = list(csv.reader(f, delimiter=sep))
                    text = "\n".join([" | ".join(row) for row in rows])
                except Exception:
                    text = self._read_text_with_encoding(abs_path)
            elif ext == "pdf":
                # Используем новый модуль pdf_reader вместо дублирования кода
                from ..pdf_reader import PdfReader
                
                reader = PdfReader()
                result = reader.read_pdf(
                    path=abs_path,
                    ocr='auto',
                    lang='rus+eng',
                    budget_seconds=5,
                    max_pages_text=100,
                    max_pages_ocr=2
                )
                
                text = result['text']
                ocr_used = result['ocr_used']
            elif ext == "docx":
                text = self._extract_docx(abs_path)
                # OCR images inside DOCX (best effort)
                ocr_text = self._ocr_docx_images(abs_path)
                if ocr_text:
                    text = (text + "\n" + ocr_text).strip()
                    ocr_used = True
            elif ext == "doc":
                text = self._extract_doc(abs_path)
            elif ext == "xlsx":
                text = self._extract_xlsx(abs_path)
            elif ext == "xls":
                text = self._extract_xls(abs_path)
        except Exception:
            self._log.exception("Ошибка извлечения текста: %s", rel_path)
            text = ""

        text = self._normalize_text(text)
        meta = {
            "format": ext.upper(),
            "length": len(text),
            "date": datetime.fromtimestamp(os.path.getmtime(abs_path)).strftime("%Y-%m-%d %H:%M"),
            "ocr": ocr_used,
            "quality": self._estimate_quality(text),
            "source": source,
        }
        return text, meta

    def _write_entry(self, out, rel_path: str, text: str, meta: dict):
        # Дополнительная защита: не пишем запись для самого индексного файла
        if os.path.basename(rel_path) == "_search_index.txt":
            return
        out.write(f"{HEADER_BAR}\n")
        out.write(f"ЗАГОЛОВОК: {rel_path}\n")
        out.write(
            "Формат: {fmt} | Символов: {length} | Дата: {date} | OCR: {ocr} | Качество: {quality}%\n".format(
                fmt=meta.get("format", ""),
                length=meta.get("length", 0),
                date=meta.get("date", ""),
                ocr="да" if meta.get("ocr") else "нет",
                quality=meta.get("quality", 0),
            )
        )
        out.write(f"Источник: {meta.get('source','filesystem')}\n")
        out.write(f"{HEADER_BAR}\n")
        out.write(f"{DOC_START_MARKER}\n")
        out.write((text or "").strip() + "\n")
        out.write(f"{DOC_END_MARKER}\n\n")

    # ---------- Helpers ----------
    def _read_text_with_encoding(self, path: str) -> str:
        try:
            import chardet  # type: ignore
        except Exception:
            chardet = None
        encodings = ["utf-8", "cp1251", "cp866", "iso-8859-1"]
        if chardet:
            try:
                with open(path, "rb") as f:
                    raw = f.read(8192)
                det = chardet.detect(raw)
                if det and det.get("encoding"):
                    encodings = [det["encoding"]] + [e for e in encodings if e != det["encoding"]]
            except Exception:
                pass
        for enc in encodings:
            try:
                with open(path, "r", encoding=enc, errors="strict") as f:
                    return f.read()
            except Exception:
                continue
        # last resort
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _normalize_text(self, text: str) -> str:
        import unicodedata, re
        if not text:
            return ""
        t = unicodedata.normalize("NFKC", text)
        t = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]", " ", t)
        t = re.sub(r"\s+", " ", t)
        return t.strip()

    def _html_to_text(self, html: str) -> str:
        """Очень простая очистка HTML без внешних зависимостей: убираем теги, скрипты/стили, декодируем сущности."""
        if not html:
            return ""
        import re, html as h
        # Удаляем содержимое <script> и <style>
        cleaned = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.IGNORECASE)
        cleaned = re.sub(r"<style[\s\S]*?</style>", " ", cleaned, flags=re.IGNORECASE)
        # Удаляем остальные теги
        cleaned = re.sub(r"<[^>]+>", " ", cleaned)
        # Декодируем HTML-сущности
        cleaned = h.unescape(cleaned)
        return cleaned

    def _estimate_quality(self, text: str) -> int:
        if not text:
            return 0
        import re
        letters = len(re.findall(r"[A-Za-zА-Яа-яЁё]", text))
        ratio = letters / max(1, len(text))
        return int(min(100, max(0, ratio * 100)))

    def _ocr_pdf_pages(self, path: str, max_pages: int = 3) -> str:
        """Best-effort OCR: конвертирует первые N страниц в изображения и распознаёт текст.
        Требует системные зависимости (poppler для pdf2image и tesseract для OCR)."""
        if max_pages <= 0:
            return ""
        try:
            from pdf2image import convert_from_path  # type: ignore
            import pytesseract  # type: ignore
            from PIL import Image  # type: ignore  # noqa: F401
        except Exception:
            # Зависимости не установлены — спокойно пропускаем
            self._log.debug("OCR пропущен: отсутствуют зависимости pdf2image/pytesseract/Pillow")
            return ""
        try:
            images = convert_from_path(path)
        except Exception:
            # Возможно, отсутствует poppler (pdftoppm/pdftocairo)
            self._log.debug("OCR пропущен: не удалось конвертировать PDF в изображения (вероятно, нет poppler)")
            return ""
        if not images:
            return ""
        texts: List[str] = []
        for img in images[:max_pages]:
            try:
                t = pytesseract.image_to_string(img, lang="rus+eng")
                if t and t.strip():
                    texts.append(t)
            except Exception:
                continue
        return "\n".join(texts).strip()

    def _extract_docx(self, path: str) -> str:
        try:
            import docx  # type: ignore
            d = docx.Document(path)
            parts = [p.text for p in d.paragraphs if p.text]
            # tables
            for table in d.tables:
                for row in table.rows:
                    row_txt = [cell.text for cell in row.cells if cell.text]
                    if row_txt:
                        parts.append(" | ".join(row_txt))
            return "\n".join(parts).strip()
        except Exception:
            return ""

    def _ocr_docx_images(self, path: str) -> str:
        # Read images from docx as zip and OCR
        try:
            import pytesseract  # type: ignore
            from PIL import Image  # type: ignore
        except Exception:
            return ""
        texts: List[str] = []
        try:
            with zipfile.ZipFile(path) as z:
                for zi in z.infolist():
                    if zi.filename.startswith("word/media/") and not zi.is_dir():
                        data = z.read(zi)
                        from io import BytesIO
                        try:
                            img = Image.open(BytesIO(data))
                            t = pytesseract.image_to_string(img, lang="rus+eng")
                            if t:
                                texts.append(t)
                        except Exception:
                            continue
        except Exception:
            return ""
        return "\n".join(texts).strip()

    def _extract_doc(self, path: str) -> str:
        # Best-effort: simple binary heuristic + cp1251 decode fallback
        try:
            with open(path, "rb") as f:
                data = f.read()
            # Keep printable ASCII and cp1251 range by replacing others with space
            filtered = bytearray()
            for b in data:
                if 32 <= b <= 126 or b in (9, 10, 13) or 192 <= b <= 255:
                    filtered.append(b)
                else:
                    filtered.append(32)
            try:
                return filtered.decode("cp1251")
            except Exception:
                return filtered.decode("latin-1", errors="ignore")
        except Exception:
            return ""

    def _extract_xlsx(self, path: str, max_rows: int = 1000, max_sheets: int = 10) -> str:
        """FR-004: Извлечение текста из XLSX с ограничениями для оптимизации."""
        try:
            import openpyxl  # type: ignore
            wb = openpyxl.load_workbook(path, data_only=True)
            out: List[str] = []
            # FR-004: Ограничиваем количество листов
            for ws in wb.worksheets[:max_sheets]:
                out.append(f"Лист: {ws.title}")
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    # FR-004: Ограничиваем количество строк на листе
                    if row_count >= max_rows:
                        break
                    vals = [str(v) for v in row if v is not None]
                    if vals:
                        out.append(" | ".join(vals))
                        row_count += 1
            return "\n".join(out).strip()
        except Exception:
            return ""

    def _extract_xls(self, path: str, max_rows: int = 1000, max_sheets: int = 10) -> str:
        """FR-004: Извлечение текста из XLS с ограничениями для оптимизации."""
        try:
            import xlrd  # type: ignore
            book = xlrd.open_workbook(path)
            out: List[str] = []
            # FR-004: Ограничиваем количество листов
            sheets_to_process = min(book.nsheets, max_sheets)
            for si in range(sheets_to_process):
                sh = book.sheet_by_index(si)
                out.append(f"Лист: {sh.name}")
                # FR-004: Ограничиваем количество строк
                rows_to_process = min(sh.nrows, max_rows)
                for r in range(rows_to_process):
                    vals = [str(sh.cell_value(r, c)) for c in range(sh.ncols) if sh.cell_value(r, c) not in ("", None)]
                    if vals:
                        out.append(" | ".join(vals))
            return "\n".join(out).strip()
        except Exception:
            return ""
