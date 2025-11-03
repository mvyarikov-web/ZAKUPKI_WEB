"""Маршруты для AI RAG анализа и связанных утилит.

Все ответы и сообщения — на русском языке. Имена сущностей в коде — на английском (PEP8).
Поддержка флагов force_web_search и clear_document_context в /ai_rag/analyze.
"""
import os
import io
import re
import json
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from flask import Blueprint, current_app, jsonify, render_template, request, send_file
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from webapp.services.rag_service import get_rag_service
from utils.api_keys_manager_multiple import get_api_keys_manager_multiple
from utils.token_tracker import (
    log_token_usage,
    get_token_stats,
    get_current_month_stats,
    get_all_time_stats,
)

ai_rag_bp = Blueprint("ai_rag", __name__, url_prefix="/ai_rag")


# ==========================
# models.json helpers + cost
# ==========================

def _models_config_path() -> str:
    try:
        root = current_app.root_path if current_app else os.getcwd()
    except Exception:
        root = os.getcwd()
    # Поддержка кастомного пути из конфига тестов
    override = None
    try:
        override = current_app.config.get("RAG_MODELS_FILE") if current_app else None
    except Exception:
        override = None
    if override:
        return override
    return os.path.join(root, "index", "models.json")


def _load_models_config_with_migration() -> Tuple[Dict[str, Any], bool]:
    """Загрузить конфиг моделей, при необходимости выполнить миграцию формата.

    Возвращает (cfg, migrated_flag).
    """
    path = _models_config_path()
    migrated = False
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                # Старый формат: просто массив моделей
                data = {
                    "models": data,
                    "default_model": data[0]["model_id"] if data else None,
                }
                migrated = True
            data.setdefault("models", [])
            if "default_model" not in data:
                data["default_model"] = data["models"][0]["model_id"] if data["models"] else None
                migrated = True
            # Приводим ключи цены из старого вида к новому
            for m in data["models"]:
                if "price_input_per_1M" in m and "price_input_per_1m" not in m:
                    m["price_input_per_1m"] = m.pop("price_input_per_1M")
                    migrated = True
                if "price_output_per_1M" in m and "price_output_per_1m" not in m:
                    m["price_output_per_1m"] = m.pop("price_output_per_1M")
                    migrated = True
            return data, migrated
    except Exception:
        current_app.logger.exception("Ошибка чтения index/models.json") if current_app else None
    # Дефолтная конфигурация
    return (
        {
            "models": [
                {
                    "model_id": "gpt-4o-mini",
                    "display_name": "GPT-4o Mini",
                    "context_window_tokens": 128000,
                    "price_input_per_1m": 0.15,
                    "price_output_per_1m": 0.60,
                    "enabled": True,
                    "supports_system_role": True,
                }
            ],
            "default_model": "gpt-4o-mini",
        },
        False,
    )


def _load_models_config() -> Dict[str, Any]:
    cfg, _ = _load_models_config_with_migration()
    return cfg


def _save_models_config(cfg: Dict[str, Any]) -> None:
    path = _models_config_path()
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)


def _calculate_cost(model_config: Dict[str, Any], usage: Dict[str, int], request_count: int = 1) -> Dict[str, Any]:
    pricing = model_config.get("pricing_model", "per_token")
    if pricing == "per_request":
        price_per_1000 = float(model_config.get("price_per_1000_requests", 0.0) or 0.0)
        per_request = price_per_1000 / 1000.0 if price_per_1000 > 0 else 0.0
        total = round(per_request * max(1, int(request_count)), 6)
        return {
            "input": 0.0,
            "output": 0.0,
            "total": total,
            "currency": "USD",
            "pricing_model": "per_request",
            "requests_count": max(1, int(request_count)),
        }
    in_toks = int(usage.get("input_tokens", 0) or 0)
    out_toks = int(usage.get("output_tokens", 0) or 0)
    price_in = float(model_config.get("price_input_per_1m", 0.0) or 0.0)
    price_out = float(model_config.get("price_output_per_1m", 0.0) or 0.0)
    inp_cost = (in_toks / 1_000_000.0) * price_in
    out_cost = (out_toks / 1_000_000.0) * price_out
    total = round(inp_cost + out_cost, 6)
    return {
        "input": round(inp_cost, 6),
        "output": round(out_cost, 6),
        "total": total,
        "currency": "USD",
        "pricing_model": "per_token",
    }


# ===========================
# Provider-aware API client
# ===========================

def _get_api_client(model_id: str, default_api_key: Optional[str], timeout: int):
    """Вернуть OpenAI-совместимый клиент с учётом провайдера (OpenAI/Perplexity/DeepSeek)."""
    try:
        from openai import OpenAI  # type: ignore
    except Exception:
        import openai  # type: ignore
        OpenAI = getattr(openai, "OpenAI", None)
        if OpenAI is None:
            raise RuntimeError("Не установлен пакет openai>=1.x (нужен класс OpenAI)")

    api_keys_mgr = None
    try:
        api_keys_mgr = get_api_keys_manager_multiple()
    except Exception:
        api_keys_mgr = None

    base_url = None
    api_key = default_api_key

    if model_id.startswith("sonar"):
        base_url = "https://api.perplexity.ai"
        if api_keys_mgr:
            api_key = api_keys_mgr.get_key("perplexity") or api_key
        api_key = api_key or os.environ.get("PPLX_API_KEY") or os.environ.get("PERPLEXITY_API_KEY") or os.environ.get("OPENAI_API_KEY")
    elif model_id.startswith("deepseek"):
        base_url = "https://api.deepseek.com"
        if api_keys_mgr:
            api_key = api_keys_mgr.get_key("deepseek") or api_key
        api_key = api_key or os.environ.get("DEEPSEEK_API_KEY") or os.environ.get("OPENAI_API_KEY")
    else:
        if api_keys_mgr:
            api_key = api_keys_mgr.get_key("openai") or api_key
        api_key = api_key or os.environ.get("OPENAI_API_KEY")

    if not api_key:
        raise RuntimeError("API ключ не настроен для выбранного провайдера")

    if base_url:
        return OpenAI(api_key=api_key, base_url=base_url, timeout=timeout)
    return OpenAI(api_key=api_key, timeout=timeout)


# ===========================
# Direct analyze (no RAG)
# ===========================

def _direct_analyze_without_rag(
    file_paths: List[str],
    prompt: str,
    model_id: str,
    max_output_tokens: int,
    temperature: float,
    upload_folder: str,
    usd_rub_rate: float = None,
    search_enabled: bool = False,
    search_params: Optional[Dict[str, Any]] = None,
    suppress_documents: bool = False,
    force_web_search: bool = False,
):
    """Прямой анализ без RAG. Поддержка подавления документов и форс-веб-поиска."""
    import time

    start_time = time.time()
    try:
        # API ключ по умолчанию (уточняется в _get_api_client)
        api_keys_mgr = get_api_keys_manager_multiple()
        default_api_key = (api_keys_mgr.get_key("openai") if api_keys_mgr else None) or current_app.config.get("OPENAI_API_KEY") or os.environ.get("OPENAI_API_KEY")

        # Тексты документов из индекса (если не подавляем)
        from webapp.routes.ai_analysis import _extract_text_from_index_for_files
        from webapp.services.indexing import get_index_path
        index_folder = current_app.config.get("INDEX_FOLDER")
        index_path = get_index_path(index_folder) if index_folder else None

        combined_docs = ""
        if not suppress_documents and index_path and os.path.exists(index_path):
            combined_docs = _extract_text_from_index_for_files(file_paths, index_path)

        if (not combined_docs or not combined_docs.strip()) and not suppress_documents:
            return jsonify({"success": False, "message": "Не удалось извлечь текст из файлов. Постройте индекс через кнопку «Построить индекс»."}), 400

        cfg = _load_models_config()
        model_config = None
        for m in cfg.get("models", []):
            if m.get("model_id") == model_id:
                model_config = m
                break
        supports_system = True if not model_config else model_config.get("supports_system_role", True)
        timeout = model_config["timeout"] if model_config and "timeout" in model_config else current_app.config.get("OPENAI_TIMEOUT", 90)

        def _get_context_window_tokens(mid: str, cfg_m: Optional[Dict[str, Any]]) -> int:
            if cfg_m:
                try:
                    val = int(cfg_m.get("context_window_tokens") or 0)
                    if val > 0:
                        return val
                except Exception:
                    pass
            known: Dict[str, int] = {
                "gpt-3.5-turbo": 16385,
                "gpt-4o-mini": 128000,
                "gpt-4o": 128000,
                "gpt-4.1": 128000,
                "gpt-5": 200000,
                "deepseek-chat": 65536,
                "deepseek-reasoner": 65536,
                "sonar": 128000,
                "sonar-pro": 128000,
                "sonar-reasoning": 128000,
                "sonar-reasoning-pro": 128000,
                "sonar-deep-research": 128000,
            }
            for k, v in known.items():
                if mid.startswith(k):
                    return v
            return 16385

        context_window = _get_context_window_tokens(model_id, model_config)

        # Оценка токенов
        try:
            import tiktoken  # type: ignore

            encoding = tiktoken.get_encoding("cl100k_base")

            def count_tokens(text: str) -> int:
                return len(encoding.encode(text))

            def truncate_by_tokens(text: str, max_tokens: int) -> str:
                if max_tokens <= 0:
                    return ""
                ids = encoding.encode(text)
                if len(ids) <= max_tokens:
                    return text
                return encoding.decode(ids[:max_tokens])
        except Exception:
            avg_chars_per_token = 4

            def count_tokens(text: str) -> int:
                return max(1, len(text) // avg_chars_per_token)

            def truncate_by_tokens(text: str, max_tokens: int) -> str:
                if max_tokens <= 0:
                    return ""
                return text[: max_tokens * avg_chars_per_token]

        reserve_for_output = int(max_output_tokens or 0)
        safety_margin = 512

        system_base = "Вы - помощник для анализа документов. Отвечайте на русском языке."
        if force_web_search:
            system_base += " Если доступен веб‑поиск, сначала выполни поиск в интернете по запросу и приведи ссылки на источники."
        if suppress_documents:
            system_base += " Игнорируй загруженные документы и основывайся на запросе пользователя и результатах веб‑поиска."

        overhead_text = (
            system_base + f"\n\nЗапрос: {prompt}"
            if supports_system
            else f"Ты - помощник для анализа документов. Отвечай на русском языке.\n\nЗапрос: {prompt}"
        )
        overhead_tokens = count_tokens(overhead_text)

        if suppress_documents:
            allowed_doc_tokens = 0
            combined_docs_trunc = ""
        else:
            allowed_doc_tokens = context_window - reserve_for_output - overhead_tokens - safety_margin
            if allowed_doc_tokens <= 0:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": (
                                f"Слишком большой объём текста для выбранной модели {model_id}. "
                                f"Доступно токенов под документы: {max(0, allowed_doc_tokens)} из {context_window}. "
                                "Уменьшите число выбранных файлов или выберите модель с большим контекстом."
                            ),
                        }
                    ),
                    400,
                )
            combined_docs_trunc = truncate_by_tokens(combined_docs, allowed_doc_tokens)

        client = _get_api_client(model_id, default_api_key, timeout)

        if supports_system:
            messages = (
                [
                    {"role": "system", "content": system_base},
                    {"role": "user", "content": f"Запрос: {prompt}"},
                ]
                if suppress_documents
                else [
                    {"role": "system", "content": system_base},
                    {"role": "user", "content": f"Документы:\n\n{combined_docs_trunc}\n\nЗапрос: {prompt}"},
                ]
            )
        else:
            messages = (
                [
                    {
                        "role": "user",
                        "content": "Ты - помощник для анализа документов. Отвечай на русском языке.\n\nЗапрос: "
                        + prompt,
                    }
                ]
                if suppress_documents
                else [
                    {
                        "role": "user",
                        "content": "Ты - помощник для анализа документов. Отвечай на русском языке.\n\nДокументы:\n\n"
                        + combined_docs_trunc
                        + "\n\nЗапрос: "
                        + prompt,
                    }
                ]
            )

        is_new_family = model_id.startswith(("o1", "o3", "o4", "gpt-4.1", "gpt-5"))
        is_sonar = model_id.startswith("sonar")

        try:
            if is_new_family:
                response = client.chat.completions.create(
                    model=model_id,
                    messages=messages,
                    max_completion_tokens=max_output_tokens,
                )
            else:
                req_kwargs: Dict[str, Any] = {"model": model_id, "messages": messages, "temperature": temperature}
                if is_sonar and (search_enabled or force_web_search):
                    try:
                        from webapp.services.search.manager import (
                            normalize_search_params,
                            apply_search_to_request,
                        )

                        norm = (
                            normalize_search_params(search_params)
                            if search_params
                            else {
                                "search_context_size": "medium",
                                "language_preference": "ru",
                            }
                        )
                        # При force_web_search принудительно устанавливаем medium для стабильного поиска
                        if force_web_search and norm:
                            norm["search_context_size"] = "medium"
                        apply_search_to_request(req_kwargs, norm or {})
                    except Exception:
                        req_kwargs["extra_body"] = {
                            "enable_search_classifier": True,
                            "search_mode": "web",
                            "language_preference": "ru",
                            "web_search_options": {"search_context_size": "medium"},
                        }
                        if search_params:
                            req_kwargs["extra_body"].update(search_params)
                else:
                    if is_sonar:
                        req_kwargs["extra_body"] = {"disable_search": True}
                    else:
                        req_kwargs["max_tokens"] = max_output_tokens

                response = client.chat.completions.create(**req_kwargs)
        except Exception as api_err:
            error_str = str(api_err)
            current_app.logger.error(f"Ошибка API {model_id}: {error_str}", exc_info=True)
            if is_sonar and (
                "401" in error_str
                or "Authorization Required" in error_str
                or "AuthenticationError" in error_str
                or "unauthorized" in error_str.lower()
            ):
                return jsonify({
                    "success": False,
                    "message": (
                        "Не удалось выполнить запрос к Perplexity (sonar): неверный или просроченный API ключ. "
                        "Проверьте ключ в разделе «API ключи» или задайте переменную окружения PPLX_API_KEY/PERPLEXITY_API_KEY."
                    ),
                }), 401
            if "timed out" in error_str.lower() or "timeout" in error_str.lower():
                return jsonify({"success": False, "message": f"Запрос превышает лимит времени {timeout} сек. Попробуйте позже."}), 504
            if "context_length_exceeded" in error_str or "maximum context length" in error_str:
                reduced_tokens = max(1, allowed_doc_tokens // 2)
                combined_docs_short = truncate_by_tokens(combined_docs_trunc, reduced_tokens)
                if supports_system:
                    messages = [
                        {"role": "system", "content": "Вы - помощник для анализа документов. Отвечайте на русском языке."},
                        {"role": "user", "content": f"Документы (сокращённые):\n\n{combined_docs_short}\n\nЗапрос: {prompt}"},
                    ]
                else:
                    messages = [
                        {"role": "user", "content": f"Ты - помощник для анализа документов. Отвечай на русском языке.\n\nДокументы (сокращённые):\n\n{combined_docs_short}\n\nЗапрос: {prompt}"},
                    ]
                try:
                    if is_new_family:
                        response = client.chat.completions.create(
                            model=model_id, messages=messages, max_completion_tokens=max_output_tokens
                        )
                    else:
                        req_kwargs = {"model": model_id, "messages": messages, "temperature": temperature}
                        if is_sonar and (search_enabled or force_web_search):
                            try:
                                from webapp.services.search.manager import (
                                    normalize_search_params,
                                    apply_search_to_request,
                                )

                                norm = normalize_search_params(search_params) if search_params else {}
                                apply_search_to_request(req_kwargs, norm or {})
                            except Exception:
                                req_kwargs["extra_body"] = {
                                    "enable_search_classifier": True,
                                    "search_mode": "web",
                                    "language_preference": "ru",
                                    "web_search_options": {"search_context_size": "medium"},
                                }
                                if search_params:
                                    req_kwargs["extra_body"].update(search_params)
                        else:
                            if is_sonar:
                                req_kwargs["extra_body"] = {"disable_search": True}
                            else:
                                req_kwargs["max_tokens"] = max_output_tokens
                        response = client.chat.completions.create(**req_kwargs)
                except Exception as retry_err:
                    return jsonify({"success": False, "message": f"Ошибка анализа даже после сокращения текста: {str(retry_err)}"}), 500
            elif (
                "rate_limit_exceeded" in error_str
                or "tokens per min" in error_str
                or "TPM" in error_str
                or "Error code: 429" in error_str
            ):
                try:
                    import time as _time

                    doc_tokens_now = count_tokens(combined_docs_trunc)
                    total_requested = overhead_tokens + doc_tokens_now + reserve_for_output
                    tpm_limit = None
                    m = re.search(r"Limit\s+(\d+)\D+Requested\s+(\d+)", error_str)
                    if m:
                        try:
                            tpm_limit = int(m.group(1))
                            requested_val = int(m.group(2))
                            current_app.logger.warning(
                                f"TPM ограничение: limit={tpm_limit}, requested={requested_val}"
                            )
                        except Exception:
                            tpm_limit = None
                    if tpm_limit and tpm_limit > 1000:
                        target_total = max(1000, int(tpm_limit * 0.9) - safety_margin)
                    else:
                        target_total = max(1000, int(total_requested * 0.5))
                    if target_total <= overhead_tokens + 64:
                        return jsonify({
                            "success": False,
                            "message": (
                                "Ограничение провайдера по токенам в минуту для модели слишком низкое для текущего объёма. "
                                "Уменьшите количество файлов, сократите промпт или попробуйте позже/выберите другую модель."
                            ),
                        }), 429
                    ratio = target_total / max(1, total_requested)
                    new_doc_token_limit = max(1, int(doc_tokens_now * ratio))
                    new_output_limit = max(64, int(reserve_for_output * ratio))
                    _time.sleep(1)
                    if is_new_family:
                        response = client.chat.completions.create(
                            model=model_id, messages=messages, max_completion_tokens=new_output_limit
                        )
                    else:
                        req_kwargs = {"model": model_id, "messages": messages, "temperature": temperature}
                        if is_sonar and (search_enabled or force_web_search):
                            try:
                                from webapp.services.search.manager import (
                                    normalize_search_params,
                                    apply_search_to_request,
                                )

                                norm = normalize_search_params(search_params) if search_params else {}
                                apply_search_to_request(req_kwargs, norm or {})
                            except Exception:
                                req_kwargs["extra_body"] = {
                                    "enable_search_classifier": True,
                                    "search_mode": "web",
                                    "language_preference": "ru",
                                    "web_search_options": {"search_context_size": "medium"},
                                }
                                if search_params:
                                    req_kwargs["extra_body"].update(search_params)
                        else:
                            if is_sonar:
                                req_kwargs["extra_body"] = {"disable_search": True}
                            else:
                                req_kwargs["max_tokens"] = new_output_limit
                        response = client.chat.completions.create(**req_kwargs)
                except Exception as retry_rate_err:
                    return jsonify({
                        "success": False,
                        "message": (
                            "Провайдер вернул ограничение по токенам в минуту (429). Даже после автоматического уменьшения объёма запрос не прошёл. "
                            f"Детали: {str(retry_rate_err)}"
                        ),
                    }), 429
            else:
                return jsonify({"success": False, "message": f"Ошибка анализа: {error_str}"}), 500

        answer = response.choices[0].message.content
        finish_reason = response.choices[0].finish_reason
        if finish_reason == "length":
            current_app.logger.warning(
                f"Ответ был обрезан из-за лимита max_tokens={max_output_tokens}"
            )
            answer += (
                "\n\n⚠️ Примечание: Ответ был обрезан из-за ограничения длины. Увеличьте параметр max_output_tokens."
            )
        usage = {
            "input_tokens": response.usage.prompt_tokens,
            "output_tokens": response.usage.completion_tokens,
            "total_tokens": response.usage.total_tokens,
        }
        duration_seconds = time.time() - start_time
        log_token_usage(
            model_id=model_id,
            prompt_tokens=usage["input_tokens"],
            completion_tokens=usage["output_tokens"],
            total_tokens=usage["total_tokens"],
            duration_seconds=duration_seconds,
            metadata={
                "file_count": len(file_paths),
                "prompt_length": len(prompt),
                "mode": "direct_without_rag",
            },
        )
        model_display = model_id
        if is_sonar and (search_enabled or force_web_search):
            model_display = f"{model_id} + Search"
        result = {
            "answer": answer,
            "usage": usage,
            "model": model_display,
            "finish_reason": finish_reason,
        }
        if model_config:
            cost = _calculate_cost(model_config, usage, request_count=1)
            usd_to_rub = (
                usd_rub_rate if (usd_rub_rate and usd_rub_rate > 0) else current_app.config.get("USD_TO_RUB_RATE", 95.0)
            )
            result["cost"] = {
                "input": cost["input"],
                "output": cost["output"],
                "total": cost["total"],
                "currency": cost["currency"],
                "pricing_model": cost.get("pricing_model", "per_token"),
                "input_rub": round(cost["input"] * usd_to_rub, 2),
                "output_rub": round(cost["output"] * usd_to_rub, 2),
                "total_rub": round(cost["total"] * usd_to_rub, 2),
                "usd_to_rub_rate": usd_to_rub,
            }
            if cost.get("pricing_model") == "per_request":
                result["cost"]["requests_count"] = cost.get("requests_count", 1)
        return jsonify({"success": True, "message": "Анализ выполнен успешно (без RAG)", "result": result}), 200

    except Exception as e:
        current_app.logger.exception(f"Ошибка прямого анализа: {e}")
        return jsonify({"success": False, "message": f"Ошибка анализа: {str(e)}"}), 500


# ===========================
# HTML render endpoint
# ===========================

@ai_rag_bp.route("/render_html", methods=["POST"])
def render_html():
    """Рендер результата анализа в HTML."""
    try:
        from webapp.utils.markdown_renderer import render_analysis_result

        data = request.get_json(silent=True) or {}
        result = data.get("result")
        if not result:
            return jsonify({"success": False, "message": "Не передан результат для рендеринга"}), 400
        html = render_analysis_result(result)
        return jsonify({"success": True, "html": html}), 200
    except Exception as e:
        current_app.logger.exception(f"Ошибка рендеринга HTML: {e}")
        return jsonify({"success": False, "message": f"Ошибка рендеринга: {str(e)}"}), 500


# ===========================
# Analyze endpoint
# ===========================

@ai_rag_bp.route("/analyze", methods=["POST"])
def analyze():
    """Выполнить анализ: RAG или прямой путь с опциями."""
    import time

    start_time = time.time()
    try:
        data = request.get_json(silent=True) or {}
        file_paths = data.get("file_paths", [])
        prompt = data.get("prompt", "")
        model_id = data.get("model_id", "gpt-4o-mini")
        top_k = data.get("top_k", 5)
        max_output_tokens = data.get("max_output_tokens", 600)
        temperature = data.get("temperature", 0.3)
        usd_rub_rate = data.get("usd_rub_rate")
        force_web_search = bool(data.get("force_web_search", False))
        clear_document_context = bool(data.get("clear_document_context", False))
        search_enabled = bool(data.get("search_enabled", False))
        search_params = data.get("search_params", {}) if search_enabled else {}

        # тотальное логирование входа
        try:
            _save_ai_analyze_artifacts(
                kind="request",
                payload={
                    "file_paths": file_paths,
                    "prompt": prompt,
                    "model_id": model_id,
                    "top_k": top_k,
                    "max_output_tokens": max_output_tokens,
                    "temperature": temperature,
                    "extra": {k: v for k, v in data.items() if k not in {"file_paths", "prompt", "model_id", "top_k", "max_output_tokens", "temperature"}},
                },
            )
        except Exception:
            current_app.logger.debug("Не удалось сохранить last_ai_analyze_request.json", exc_info=True)

        if not file_paths:
            return jsonify({"success": False, "message": "Не выбраны файлы для анализа"}), 400
        if not prompt:
            return jsonify({"success": False, "message": "Не указан промпт"}), 400

        upload_folder = current_app.config["UPLOAD_FOLDER"]
        rag_service = get_rag_service()

        # Web-only режим или очистка контекста — идём в прямой путь без документов
        if force_web_search or clear_document_context:
            current_app.logger.info("Пропускаем RAG: force_web_search/clear_document_context активны")
            resp = _direct_analyze_without_rag(
                file_paths=file_paths,
                prompt=prompt,
                model_id=model_id,
                max_output_tokens=max_output_tokens,
                temperature=temperature,
                upload_folder=upload_folder,
                usd_rub_rate=usd_rub_rate,
                search_enabled=search_enabled,
                search_params=search_params,
                suppress_documents=True,
                force_web_search=True,
            )
            try:
                payload_to_save = None
                if isinstance(resp, tuple) and len(resp) >= 2 and hasattr(resp[0], "get_json"):
                    payload_to_save = resp[0].get_json(silent=True)
                _save_ai_analyze_artifacts(kind="response", payload=payload_to_save)
            except Exception:
                current_app.logger.debug(
                    "Не удалось сохранить last_ai_analyze_result.json (web-only path)",
                    exc_info=True,
                )
            return resp

        # Если база недоступна — fallback напрямую
        if not rag_service.db_available:
            current_app.logger.warning("БД недоступна, используется прямой анализ без RAG")
            resp = _direct_analyze_without_rag(
                file_paths=file_paths,
                prompt=prompt,
                model_id=model_id,
                max_output_tokens=max_output_tokens,
                temperature=temperature,
                upload_folder=upload_folder,
                usd_rub_rate=usd_rub_rate,
                search_enabled=search_enabled,
                search_params=search_params,
            )
            try:
                payload_to_save = None
                if isinstance(resp, tuple) and len(resp) >= 2 and hasattr(resp[0], "get_json"):
                    payload_to_save = resp[0].get_json(silent=True)
                _save_ai_analyze_artifacts(kind="response", payload=payload_to_save)
            except Exception:
                current_app.logger.debug(
                    "Не удалось сохранить last_ai_analyze_result.json (fallback path)",
                    exc_info=True,
                )
            return resp

        # Основной путь: RAG
        try:
            success, message, result = rag_service.search_and_analyze(
                query=prompt,
                file_paths=file_paths,
                model=model_id,
                top_k=top_k,
                max_output_tokens=max_output_tokens,
                temperature=temperature,
                upload_folder=upload_folder,
                search_params=search_params if search_enabled else None,
            )
        except Exception as rag_err:
            current_app.logger.warning(f"Ошибка RAG, переключение на прямой анализ: {rag_err}")
            resp = _direct_analyze_without_rag(
                file_paths=file_paths,
                prompt=prompt,
                model_id=model_id,
                max_output_tokens=max_output_tokens,
                temperature=temperature,
                upload_folder=upload_folder,
                usd_rub_rate=usd_rub_rate,
                search_enabled=search_enabled,
                search_params=search_params,
            )
            try:
                payload_to_save = None
                if isinstance(resp, tuple) and len(resp) >= 2 and hasattr(resp[0], "get_json"):
                    payload_to_save = resp[0].get_json(silent=True)
                _save_ai_analyze_artifacts(kind="response", payload=payload_to_save)
            except Exception:
                current_app.logger.debug(
                    "Не удалось сохранить last_ai_analyze_result.json (rag error path)",
                    exc_info=True,
                )
            return resp

        if not success:
            # для ошибок RAG можно fallback
            if any(
                kw in message.lower()
                for kw in ["эмбеддинг", "embedding", "база данных недоступна", "database"]
            ):
                current_app.logger.warning(
                    f"Ошибка RAG/БД, переключение на прямой анализ: {message}"
                )
                resp = _direct_analyze_without_rag(
                    file_paths=file_paths,
                    prompt=prompt,
                    model_id=model_id,
                    max_output_tokens=max_output_tokens,
                    temperature=temperature,
                    upload_folder=upload_folder,
                    usd_rub_rate=usd_rub_rate,
                    search_enabled=search_enabled,
                    search_params=search_params,
                )
                try:
                    payload_to_save = None
                    if isinstance(resp, tuple) and len(resp) >= 2 and hasattr(resp[0], "get_json"):
                        payload_to_save = resp[0].get_json(silent=True)
                    _save_ai_analyze_artifacts(kind="response", payload=payload_to_save)
                except Exception:
                    current_app.logger.debug(
                        "Не удалось сохранить last_ai_analyze_result.json (rag fail path)",
                        exc_info=True,
                    )
                return resp
            return jsonify({"success": False, "message": message}), 400

        # success=True
        duration_seconds = time.time() - start_time
        usage = result.get("usage", {}) or {}
        if usage:
            log_token_usage(
                model_id=model_id,
                prompt_tokens=usage.get("input_tokens", 0),
                completion_tokens=usage.get("output_tokens", 0),
                total_tokens=usage.get("total_tokens", 0),
                duration_seconds=duration_seconds,
                metadata={
                    "file_count": len(file_paths),
                    "top_k": top_k,
                    "prompt_length": len(prompt),
                },
            )

        # стоимость
        config = _load_models_config()
        model_config = None
        for m in config.get("models", []):
            if m.get("model_id") == model_id:
                model_config = m
                break
        if model_config:
            cost = _calculate_cost(model_config, usage, request_count=1)
            usd_to_rub = (
                usd_rub_rate if (usd_rub_rate and usd_rub_rate > 0) else current_app.config.get("USD_TO_RUB_RATE", 95.0)
            )
            result["cost"] = {
                "input": cost["input"],
                "output": cost["output"],
                "total": cost["total"],
                "currency": cost["currency"],
                "pricing_model": cost.get("pricing_model", "per_token"),
                "input_rub": round(cost["input"] * usd_to_rub, 2),
                "output_rub": round(cost["output"] * usd_to_rub, 2),
                "total_rub": round(cost["total"] * usd_to_rub, 2),
                "usd_to_rub_rate": usd_to_rub,
            }
            if cost.get("pricing_model") == "per_request":
                result["cost"]["requests_count"] = cost.get("requests_count", 1)

        resp = jsonify({"success": True, "message": result.get("message", "Анализ выполнен"), "result": result}), 200
        try:
            payload_to_save = resp[0].get_json(silent=True) if hasattr(resp[0], "get_json") else None
            _save_ai_analyze_artifacts(kind="response", payload=payload_to_save)
        except Exception:
            current_app.logger.debug(
                "Не удалось сохранить last_ai_analyze_result.json (rag success)",
                exc_info=True,
            )
        return resp

    except Exception as e:
        current_app.logger.exception(f"Ошибка в /ai_rag/analyze: {e}")
        resp = jsonify({"success": False, "message": f"Внутренняя ошибка сервера: {str(e)}"}), 500
        try:
            payload_to_save = resp[0].get_json(silent=True) if hasattr(resp[0], "get_json") else None
            _save_ai_analyze_artifacts(kind="response", payload=payload_to_save)
        except Exception:
            current_app.logger.debug(
                "Не удалось сохранить last_ai_analyze_result.json (exception path)",
                exc_info=True,
            )
        return resp


# ===========================
# Status endpoint
# ===========================

@ai_rag_bp.route("/status", methods=["GET"])
def get_status():
    try:
        rag_service = get_rag_service()
        db_available = rag_service.db_available
        db_stats = None
        if db_available:
            try:
                db_stats = rag_service.get_database_stats()
            except Exception:
                pass
        api_key_available = bool(rag_service.api_key)
        return jsonify({
            "success": True,
            "rag_enabled": current_app.config.get("RAG_ENABLED", False),
            "database_available": db_available,
            "database_stats": db_stats,
            "api_key_configured": api_key_available,
            "embeddings_model": current_app.config.get("RAG_EMBEDDING_MODEL", "text-embedding-3-small"),
        }), 200
    except Exception as e:
        current_app.logger.exception(f"Ошибка в /ai_rag/status: {e}")
        return jsonify({"success": False, "message": f"Ошибка получения статуса: {str(e)}"}), 500


# ===========================
# Export DOCX endpoint
# ===========================

@ai_rag_bp.route("/export_docx", methods=["POST"])
def export_docx():
    try:
        data = request.get_json(silent=True) or {}
        result = data.get("result") or {}
        answer = result.get("answer", "")
        model = result.get("model", "неизвестная модель")
        usage = result.get("usage", {})
        cost = result.get("cost", {})
        if not answer:
            return jsonify({"success": False, "message": "Нет текста для экспорта"}), 400

        doc = Document()
        heading = doc.add_heading("Результат AI Анализа", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()
        meta_para = doc.add_paragraph()
        meta_para.add_run("Дата: ").bold = True
        meta_para.add_run(datetime.now().strftime("%d.%m.%Y %H:%M:%S"))
        meta_para.add_run("\n")
        meta_para.add_run("Модель: ").bold = True
        meta_para.add_run(model)
        meta_para.add_run("\n")
        if cost:
            meta_para.add_run("Стоимость: ").bold = True
            meta_para.add_run(f"${cost.get('total', 0):.6f} ")
            meta_para.add_run(
                f"(вход: ${cost.get('input', 0):.6f}, выход: ${cost.get('output', 0):.6f})"
            )
            meta_para.add_run("\n")
            if "total_rub" in cost:
                meta_para.add_run("В рублях: ").bold = True
                meta_para.add_run(f"₽{cost.get('total_rub', 0):.2f} ")
                meta_para.add_run(
                    f"(вход: ₽{cost.get('input_rub', 0):.2f}, выход: ₽{cost.get('output_rub', 0):.2f}) "
                )
                meta_para.add_run(f"по курсу ${cost.get('usd_to_rub_rate', 95.0):.2f}")
                meta_para.add_run("\n")
        if usage:
            meta_para.add_run("Токены: ").bold = True
            meta_para.add_run(f"{usage.get('total_tokens', 0)} ")
            meta_para.add_run(
                f"(вход: {usage.get('input_tokens', 0)}, выход: {usage.get('output_tokens', 0)})"
            )

        doc.add_paragraph("_" * 80)
        doc.add_paragraph()
        _add_markdown_to_docx(doc, answer)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"ai_analysis_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"
        return send_file(
            buffer,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )
    except Exception as e:
        current_app.logger.exception(f"Ошибка в /ai_rag/export_docx: {e}")
        return jsonify({"success": False, "message": f"Ошибка экспорта DOCX: {str(e)}"}), 500


# ===========================
# Markdown helpers for DOCX
# ===========================

def _add_markdown_to_docx(doc: Document, markdown_text: str) -> None:
    lines = markdown_text.split("\n")
    in_code_block = False
    code_block_lines: List[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code_block:
                if code_block_lines:
                    code_para = doc.add_paragraph("\n".join(code_block_lines))
                    code_para.style = "No Spacing"
                    for run in code_para.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.strip().startswith(("- ", "* ")):
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, line.strip()[2:])
            i += 1
            continue
        if re.match(r"^\d+\.\s", line.strip()):
            text = re.sub(r"^\d+\.\s+", "", line.strip())
            para = doc.add_paragraph(style="List Number")
            _add_formatted_text(para, text)
            i += 1
            continue
        para = doc.add_paragraph()
        _add_formatted_text(para, line)
        i += 1


def _add_formatted_text(paragraph, text: str) -> None:
    link_pattern = r"\[([^\]]+)\]\(([^\)]+)\)"
    parts = re.split(f"({link_pattern})", text)
    i = 0
    while i < len(parts):
        part = parts[i]
        link_match = re.match(link_pattern, part)
        if link_match:
            link_text = link_match.group(1)
            link_url = link_match.group(2)
            _add_hyperlink(paragraph, link_url, link_text)
            i += 1
            continue
        code_parts = re.split(r"(`[^`]+`)", part)
        for code_part in code_parts:
            if code_part.startswith("`") and code_part.endswith("`"):
                run = paragraph.add_run(code_part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(200, 0, 0)
            else:
                _add_bold_italic_text(paragraph, code_part)
        i += 1


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_bold_italic_text(paragraph, text: str) -> None:
    bold_parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for bold_part in bold_parts:
        if bold_part.startswith("**") and bold_part.endswith("**"):
            run = paragraph.add_run(bold_part[2:-2])
            run.bold = True
        else:
            italic_parts = re.split(r"(\*[^*]+\*)", bold_part)
            for italic_part in italic_parts:
                if italic_part.startswith("*") and italic_part.endswith("*") and not italic_part.startswith("**"):
                    run = paragraph.add_run(italic_part[1:-1])
                    run.italic = True
                else:
                    if italic_part:
                        paragraph.add_run(italic_part)


# ===========================
# Total logging artifacts
# ===========================

def _save_ai_analyze_artifacts(kind: str, payload: Optional[Dict[str, Any]]) -> None:
    try:
        if payload is None:
            return
        index_folder = current_app.config.get("INDEX_FOLDER")
        if not index_folder:
            return
        os.makedirs(index_folder, exist_ok=True)
        fname = (
            "last_ai_analyze_request.json" if kind == "request" else "last_ai_analyze_result.json"
        )
        path = os.path.join(index_folder, fname)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    except Exception:
        current_app.logger.debug(
            f"Не удалось сохранить артефакт analyze ({kind})", exc_info=True
        )


# ===========================
# Token stats endpoints
# ===========================

@ai_rag_bp.route("/token_report")
def token_report():
    return render_template("token_report.html")


@ai_rag_bp.route("/token_stats", methods=["GET"])
def token_stats():
    try:
        period = request.args.get("period", "all_time")
        start_date = request.args.get("start_date")
        end_date = request.args.get("end_date")
        model_id = request.args.get("model_id")
        if period == "current_month":
            stats = get_current_month_stats()
        elif period == "all_time":
            stats = get_all_time_stats()
        else:
            stats = get_token_stats(start_date=start_date, end_date=end_date, model_id=model_id)
        return jsonify(stats), 200
    except Exception as e:
        current_app.logger.exception(f"Ошибка получения статистики токенов: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


# ===========================
# Models endpoints
# ===========================

@ai_rag_bp.route("/models", methods=["GET"])
def get_models():
    try:
        cfg, migrated = _load_models_config_with_migration()
        if migrated:
            try:
                _save_models_config(cfg)
            except Exception:
                current_app.logger.debug("Не удалось сохранить миграцию models.json", exc_info=True)
        return jsonify({"success": True, **cfg}), 200
    except Exception as e:
        current_app.logger.exception("Ошибка чтения моделей")
        return jsonify({"success": False, "message": f"Ошибка: {str(e)}"}), 500


@ai_rag_bp.route("/models", methods=["POST"])
def update_model_prices():
    try:
        data = request.get_json(silent=True) or {}
        model_id = data.get("model_id")
        if not model_id:
            return jsonify({"success": False, "message": "Не указан model_id"}), 400
        cfg = _load_models_config()
        models = cfg.setdefault("models", [])
        entry = None
        for m in models:
            if m.get("model_id") == model_id:
                entry = m
                break
        if entry is None:
            entry = {"model_id": model_id}
            models.append(entry)
        updatable = {
            "price_input_per_1m",
            "price_output_per_1m",
            "pricing_model",
            "price_per_1000_requests",
            "timeout",
            "context_window_tokens",
            "supports_system_role",
            "provider",
        }
        for k in updatable:
            if k in data:
                entry[k] = data[k]
        _save_models_config(cfg)
        return jsonify({"success": True, "message": "Параметры модели обновлены", "model": entry}), 200
    except Exception as e:
        current_app.logger.exception("Ошибка обновления модели")
        return jsonify({"success": False, "message": f"Не удалось обновить модель: {str(e)}"}), 500


# ===========================
# Models: available/add/default/delete/search_params
# ===========================

def _infer_model_defaults(model_id: str) -> Dict[str, Any]:
    """Возвращает заготовку записи модели по её идентификатору.

    Логика определения провайдера и флагов минимальная и безопасная:
    - OpenAI: id содержит 'gpt-' или начинается с 'o1'/'o3'.
    - Perplexity: id содержит 'sonar' или 'pplx'.
    - DeepSeek: id содержит 'deepseek'.
    - supports_search: True для Perplexity (sonar...) и для id с суффиксом '-online'.
    Цены по умолчанию нулевые (редактируются в UI), для Search — цена за 1000 запросов.
    """
    mid = (model_id or "").strip()
    prov = "other"
    if "deepseek" in mid:
        prov = "deepseek"
    elif "sonar" in mid or "pplx" in mid:
        prov = "perplexity"
    elif mid.startswith("o1") or mid.startswith("o3") or "gpt-" in mid:
        prov = "openai"

    display = {
        "gpt-4o-mini": "GPT-4o mini",
        "gpt-4o": "GPT-4o",
        "o3-mini": "O3 mini",
        "o1-mini": "O1 mini",
        "deepseek-chat": "DeepSeek Chat",
        "deepseek-reasoner": "DeepSeek Reasoner",
        "sonar": "Perplexity Sonar",
        "sonar-pro": "Perplexity Sonar Pro",
    }.get(mid, mid)

    supports_search = (prov == "perplexity") or mid.endswith("-online")

    entry: Dict[str, Any] = {
        "model_id": mid,
        "display_name": display,
        "provider": prov,
        "enabled": True,
        "supports_system_role": True,
        "context_window_tokens": 128_000,
        "timeout": 30,
        # Тарификация по умолчанию: токенная, значения 0 — пользователь задаст в UI
        "pricing_model": "per_token",
        "price_input_per_1m": 0.0,
        "price_output_per_1m": 0.0,
        # Режим Search (для Perplexity и online-моделей) — выключен по умолчанию; цена по запросам
        "supports_search": supports_search,
        "search_enabled": False,
        "price_per_1000_requests": 5.0,
    }
    return entry


@ai_rag_bp.route("/models/available", methods=["GET"])
def list_available_models():
    """Возвращает список рекомендуемых доступных моделей (без сетевых вызовов).

    UI сам отфильтрует уже добавленные. Если потребуется, можно расширить список.
    """
    try:
        recommended = [
            _infer_model_defaults(x)
            for x in [
                # OpenAI
                "gpt-4o-mini",
                "gpt-4o",
                "o3-mini",
                "o1-mini",
                # DeepSeek
                "deepseek-chat",
                "deepseek-reasoner",
                # Perplexity
                "sonar",
                "sonar-pro",
            ]
        ]
        # Для ответа оставляем только ключевые поля, чтобы не засорять UI
        minimal = [
            {
                "model_id": m.get("model_id"),
                "display_name": m.get("display_name"),
                "provider": m.get("provider"),
                "context_window_tokens": m.get("context_window_tokens"),
            }
            for m in recommended
        ]
        return jsonify({"success": True, "models": minimal}), 200
    except Exception as e:
        current_app.logger.exception("Ошибка /ai_rag/models/available")
        return jsonify({"success": False, "message": f"Ошибка: {str(e)}"}), 500


@ai_rag_bp.route("/models/add", methods=["POST"])
def add_models():
    """Добавляет выбранные модели в конфигурацию models.json.

    Тело: {"model_ids": ["gpt-4o-mini", ...]}
    """
    try:
        data = request.get_json(silent=True) or {}
        model_ids = data.get("model_ids") or []
        if not isinstance(model_ids, list) or not model_ids:
            return jsonify({"success": False, "message": "Список model_ids пуст"}), 400

        cfg = _load_models_config()
        models = cfg.setdefault("models", [])
        existing_ids = {m.get("model_id") for m in models}

        added = 0
        for mid in model_ids:
            if mid in existing_ids:
                continue
            entry = _infer_model_defaults(mid)
            models.append(entry)
            added += 1

        # Если ранее не было default_model — выставим первую добавленную
        if added > 0 and not cfg.get("default_model"):
            cfg["default_model"] = model_ids[0]

        _save_models_config(cfg)
        return jsonify({"success": True, "added": added}), 200
    except Exception as e:
        current_app.logger.exception("Ошибка /ai_rag/models/add")
        return jsonify({"success": False, "message": f"Не удалось добавить модели: {str(e)}"}), 500


@ai_rag_bp.route("/models/default", methods=["PUT"])
def set_default_model():
    try:
        data = request.get_json(silent=True) or {}
        model_id = data.get("model_id")
        if not model_id:
            return jsonify({"success": False, "message": "Не указан model_id"}), 400
        cfg = _load_models_config()
        models = cfg.get("models", [])
        if not any(m.get("model_id") == model_id for m in models):
            return jsonify({"success": False, "message": "Такая модель не добавлена"}), 400
        cfg["default_model"] = model_id
        _save_models_config(cfg)
        return jsonify({"success": True, "message": "Модель по умолчанию обновлена", "default_model": model_id}), 200
    except Exception as e:
        current_app.logger.exception("Ошибка /ai_rag/models/default")
        return jsonify({"success": False, "message": f"Не удалось обновить: {str(e)}"}), 500


@ai_rag_bp.route("/models/search_params", methods=["POST"])
def save_search_params():
    try:
        data = request.get_json(silent=True) or {}
        model_id = data.get("model_id")
        params = data.get("search_params")
        if not model_id or not isinstance(params, dict):
            return jsonify({"success": False, "message": "Неверные данные"}), 400
        cfg = _load_models_config()
        models = cfg.setdefault("models", [])
        for m in models:
            if m.get("model_id") == model_id:
                m["search_params"] = params
                _save_models_config(cfg)
                return jsonify({"success": True}), 200
        return jsonify({"success": False, "message": "Модель не найдена"}), 404
    except Exception as e:
        current_app.logger.exception("Ошибка /ai_rag/models/search_params")
        return jsonify({"success": False, "message": f"Не удалось сохранить: {str(e)}"}), 500


@ai_rag_bp.route("/models/<model_id>", methods=["DELETE"])
def delete_model(model_id: str):
    try:
        cfg = _load_models_config()
        models = cfg.get("models", [])
        if not models:
            return jsonify({"success": False, "error": "Список моделей пуст"}), 400
        # Нельзя удалить последнюю модель
        if len(models) == 1 and models[0].get("model_id") == model_id:
            return jsonify({"success": False, "error": "Нельзя удалить последнюю модель"}), 400

        new_models = [m for m in models if m.get("model_id") != model_id]
        if len(new_models) == len(models):
            return jsonify({"success": False, "error": "Модель не найдена"}), 404
        cfg["models"] = new_models
        # Если удалили default_model — переназначим
        if cfg.get("default_model") == model_id:
            cfg["default_model"] = new_models[0].get("model_id") if new_models else None
        _save_models_config(cfg)
        return jsonify({"success": True, "message": "Модель удалена"}), 200
    except Exception as e:
        current_app.logger.exception("Ошибка удаления модели")
        return jsonify({"success": False, "error": str(e)}), 500
